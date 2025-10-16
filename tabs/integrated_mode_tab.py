import os
import json
import base64
import hashlib
import random
import string
import subprocess
import time
import uuid
from datetime import datetime

from PyQt5.QtWidgets import (
    QGroupBox, QComboBox, QWidget, QPushButton, QTextEdit, QVBoxLayout,
    QMessageBox, QFileDialog, QHBoxLayout, QFrame, QListWidget, QLabel,
    QTableWidget, QTableWidgetItem, QHeaderView, QProgressBar, QRadioButton,
    QScrollArea, QGridLayout, QLineEdit, QCheckBox, QSpinBox, QTabWidget,
    QPlainTextEdit, QSplitter, QButtonGroup
)
from PyQt5.QtCore import Qt, QTimer
from PyQt5.QtGui import QFont, QPixmap, QIcon

from docx import Document
import msoffcrypto

from Crypto.PublicKey import RSA
from Crypto.Cipher import AES, PKCS1_OAEP
from Crypto.Util.Padding import pad, unpad
from Crypto.Random import get_random_bytes

from cryptography.hazmat.primitives.asymmetric import padding
from cryptography.hazmat.primitives import serialization, hashes

from PIL import Image
import stepic

from pydub import AudioSegment

from mutagen.mp4 import MP4
from mutagen.id3 import ID3, TextFrame
from mutagen.flac import FLAC
from mutagen.oggvorbis import OggVorbis
from mutagen.wave import WAVE

import numpy as np
import wave

from tabs.video_tab import VideoSteganographyWorker




class AdvancedSteganoWorkflowItem:
    def __init__(self, mode_id, mode_name, config=None):
        self.mode_id = mode_id
        self.mode_name = mode_name
        self.config = config or {}
        self.source_files = []
        self.output_path = ""
        self.output_dir = ""  
        self.created_at = datetime.now()

class IntegrationTab(QWidget):
    def __init__(self):
        super().__init__()
        self.selected_files = []
        self.workflow_items = []
        self.output_path = ""
        self.current_mode_config = {}
        self.selected_extract_files = []
        self.current_extract_mode_id = 1

        self.initUI()
                
        self.on_mode_changed(0)  
        self.on_extract_mode_changed(0)           

    def initUI(self):
        
        self.setStyleSheet("""
        /* Main Widget Styling */
        QWidget {
            background: transparent;
            color: #ffffff;
            font-family: 'Segoe UI', 'Roboto', Arial, sans-serif;
        }
        
        /* Group Box Styling */
        QGroupBox {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #2d2d44, stop:1 #1e1e2e);
            border: 2px solid #00d4ff;
            border-radius: 12px;
            margin-top: 15px;
            padding: 20px;
            font-size: 14px;
            font-weight: bold;
            color: #00d4ff;
        }
        
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 20px;
            padding: 0 10px 0 10px;
            color: #00d4ff;
            font-size: 16px;
            font-weight: bold;
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #1e1e2e, stop:1 #2d2d44);
            border-radius: 8px;
        }
        
        /* Tab Widget Styling */
        QTabWidget::pane {
            border: 2px solid #00d4ff;
            border-radius: 8px;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #2d2d44, stop:1 #1e1e2e);
        }
        
        QTabBar::tab {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3a3a54, stop:1 #2a2a3e);
            color: #cccccc;
            padding: 8px 16px;
            margin-right: 2px;
            border-top-left-radius: 8px;
            border-top-right-radius: 8px;
            font-weight: bold;
        }
        
        QTabBar::tab:selected {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #00d4ff, stop:1 #0099cc);
            color: white;
        }
        
        QTabBar::tab:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4a4a64, stop:1 #3a3a4e);
        }
        
        /* Button Styling */
        QPushButton {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4a90e2, stop:1 #357abd);
            color: white;
            border: none;
            border-radius: 8px;
            padding: 12px 18px;
            font-size: 13px;
            font-weight: bold;
            min-height: 20px;
        }
        
        QPushButton:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #5ba0f2, stop:1 #4585c7);
        }
        
        QPushButton:pressed {
            background: #2d5aa0;
        }
        
        QPushButton#modeButton {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #8e24aa, stop:1 #5e35b1);
            font-size: 14px;
            padding: 15px 20px;
        }
        
        QPushButton#modeButton:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ba68c8, stop:1 #9575cd);
        }
        
        QPushButton#executeButton {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4caf50, stop:1 #388e3c);
            font-size: 15px;
            padding: 15px 25px;
        }
        
        QPushButton#executeButton:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #66bb6a, stop:1 #4caf50);
        }

        QPushButton#clearButton {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ff4444, stop:1 #cc3333);
        }
        QPushButton#clearButton:hover {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #ff6666, stop:1 #dd4444);
        }
        
        /* ComboBox Styling */
        QComboBox {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3a3a54, stop:1 #2a2a3e);
            color: #ffffff;
            border: 2px solid #555;
            border-radius: 8px;
            padding: 8px 12px;
            font-size: 13px;
            font-weight: bold;
            min-width: 120px;
        }
        
        QComboBox:hover {
            border-color: #00d4ff;
        }
        
        QComboBox:focus {
            border-color: #00d4ff;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4a4a64, stop:1 #3a3a4e);
        }
        
        QComboBox::drop-down {
            border: none;
            width: 30px;
            background: transparent;
        }
        
        QComboBox::down-arrow {
            image: none;
            border-left: 5px solid transparent;
            border-right: 5px solid transparent;
            border-top: 8px solid #00d4ff;
            margin-right: 8px;
        }
        
        QComboBox QAbstractItemView {
            background: #2a2a3e;
            border: 2px solid #00d4ff;
            border-radius: 8px;
            selection-background-color: #00d4ff;
            selection-color: white;
            padding: 6px;
        }
        
        /* LineEdit Styling */
        QLineEdit {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3a3a54, stop:1 #2a2a3e);
            color: #ffffff;
            border: 2px solid #555;
            border-radius: 8px;
            padding: 8px 12px;
            font-size: 13px;
            font-weight: bold;
        }
        
        QLineEdit:hover {
            border-color: #00d4ff;
        }
        
        QLineEdit:focus {
            border-color: #00d4ff;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #4a4a64, stop:1 #3a3a4e);
        }
        
        /* Text Edit Styling */
        QTextEdit, QPlainTextEdit {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #1e1e2e, stop:1 #2d2d44);
            color: #ffffff;
            border: 2px solid #555;
            border-radius: 8px;
            font-family: 'Consolas', 'Monaco', monospace;
            font-size: 12px;
            padding: 12px;
            line-height: 1.4;
        }
        
        QTextEdit:focus, QPlainTextEdit:focus {
            border-color: #00d4ff;
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #2e2e3e, stop:1 #3d3d54);
        }
        
        /* Label Styling */
        QLabel {
            color: #cccccc;
            font-size: 13px;
            font-weight: bold;
            background: transparent;
        }
        
        QLabel#titleLabel {
            color: #00d4ff;
            font-size: 16px;
            font-weight: bold;
        }
        
        QLabel#warningLabel {
            color: #ff9800;
            font-size: 12px;
        }
        
        /* CheckBox Styling */
        QCheckBox {
            color: #cccccc;
            font-size: 13px;
            font-weight: bold;
            spacing: 10px;
        }
        
        QCheckBox::indicator {
            width: 18px;
            height: 18px;
            border-radius: 3px;
            border: 2px solid #00d4ff;
            background-color: #2a2a3e;
        }
        
        QCheckBox::indicator:hover {
            border: 2px solid #33ddff;
        }
        
        QCheckBox::indicator:checked {
            background-color: #00d4ff;
            border: 2px solid #00d4ff;
        }
        
        /* SpinBox Styling */
        QSpinBox {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3a3a54, stop:1 #2a2a3e);
            color: #ffffff;
            border: 2px solid #555;
            border-radius: 8px;
            padding: 8px 12px;
            font-size: 13px;
            font-weight: bold;
        }
        
        QSpinBox:hover {
            border-color: #00d4ff;
        }
        
        QSpinBox:focus {
            border-color: #00d4ff;
        }
        
        /* Progress Bar Styling */
        QProgressBar {
            background: #3c3c3c;
            border: 2px solid #555;
            border-radius: 10px;
            text-align: center;
            color: #ffffff;
            font-weight: bold;
            font-size: 12px;
            height: 25px;
        }
        
        QProgressBar::chunk {
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                stop:0 #00d4ff, stop:1 #0099cc);
            border-radius: 8px;
            margin: 2px;
        }
        
        /* Table Widget Styling */
        QTableWidget {
            background: #1e1e2e;
            border: 2px solid #00d4ff;
            border-radius: 8px;
            gridline-color: #555;
            font-size: 12px;
            color: #ffffff;
        }
        
        QTableWidget::item {
            padding: 8px;
        }
        
        QTableWidget::item:selected {
            background: #00d4ff;
            color: white;
        }
        
        QHeaderView::section {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
                stop:0 #3a3a54, stop:1 #2a2a3e);
            color: #00d4ff;
            padding: 10px;
            border: 1px solid #555;
            border-bottom: 2px solid #00d4ff;
            font-weight: bold;
            font-size: 13px;
        }
        
        /* List Widget Styling */
        QListWidget {
            background: #1e1e2e;
            border: 2px solid #00d4ff;
            border-radius: 8px;
            font-size: 12px;
            color: #ffffff;
            padding: 8px;
        }
        
        QListWidget::item {
            padding: 10px;
            margin-bottom: 3px;
            border-radius: 6px;
        }
        
        QListWidget::item:hover {
            background: rgba(0, 212, 255, 0.15);
        }
        
        QListWidget::item:selected {
            background: #00d4ff;
            color: white;
        }

        /* Radio Button Styling */
        QRadioButton {
            color: #cccccc;
            font-size: 13px;
            font-weight: bold;
            spacing: 10px;
        }
        
        QRadioButton::indicator {
            width: 18px;
            height: 18px;
            border-radius: 9px;
            border: 2px solid #00d4ff;
            background-color: #2a2a3e;
        }
        
        QRadioButton::indicator:hover {
            border: 2px solid #33ddff;
        }
        
        QRadioButton::indicator:checked {
            background-color: #00d4ff;
            border: 2px solid #00d4ff;
        }
        
        /* Scroll Bar Styling */
        QScrollBar:vertical {
            background: #2a2a3e;
            width: 12px;
            border-radius: 6px;
            margin: 0;
        }
        
        QScrollBar::handle:vertical {
            background: #00d4ff;
            border-radius: 6px;
            min-height: 20px;
            margin: 2px;
        }
        
        QScrollBar::handle:vertical:hover {
            background: #33ddff;
        }
        
        QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
            height: 0px;
        }
        
        /* Splitter Styling */
        QSplitter::handle {
            background: #00d4ff;
            width: 3px;
            height: 3px;
        }
        
        QSplitter::handle:hover {
            background: #33ddff;
        }
        """)

        
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll_area.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)  
        scroll_area.setStyleSheet("""
            QScrollArea {
                border: none;
            } """)
        
        
        container_widget = QWidget()
        main_layout = QVBoxLayout(container_widget)
        main_layout.setSpacing(15)
        main_layout.setContentsMargins(15, 15, 15, 15)
        title_layout = QHBoxLayout()
        title_layout.addStretch()
        mode_group = QGroupBox("🎛️ เลือกโหมดการทำงาน (Select Operation Mode)")
        mode_layout = QVBoxLayout()
        
        # Mode dropdown with descriptions
        mode_selection_layout = QHBoxLayout()
        mode_selection_layout.addWidget(QLabel("โหมด (Mode)"))
        self.mode_dropdown = QComboBox()
        self.mode_dropdown.setMinimumWidth(400)
        
        # Define all 10 modes
        self.modes = {
            1: "🔄 โหมด 1: AES + แบ่งข้อความครึ่งหนึ่ง (ภาพ+เสียง)",
            2: "📄 โหมด 2: DOCX + RSA + Video Metadata",
            3: "🎛️ โหมด 3: AES + แบ่งข้อความ + Stego หลายไฟล์ (3 ส่วน)",
            4: "🧬 โหมด 4: AES + RSA + Metadata Stego",
            5: "🧫 โหมด 5: GPG + Metadata + EOF Embedding",
            6: "🧩 โหมด 6: AES + LSB + Metadata + Checksum",
            7: "🔄 โหมด 7: แปลงหลายชั้น + ซ่อนหลายที่",
            8: "🧠 โหมด 8: AES + GPG + หลาย Media",
            9: "🌀 โหมด 9: Nested Stego (ตุ๊กตารัสเซีย)",
            10: "🧾 โหมด 10: Split + Layered + Time-lock"
        }
        
        for mode_id, mode_name in self.modes.items():
            self.mode_dropdown.addItem(mode_name, mode_id)
        
        self.mode_dropdown.currentIndexChanged.connect(self.on_mode_changed)
        mode_selection_layout.addWidget(self.mode_dropdown)
        mode_selection_layout.addStretch()
        
        # Mode description
        self.mode_description = QLabel()
        self.mode_description.setWordWrap(True)
        self.mode_description.setStyleSheet("color: #ffeb3b; font-size: 12px; padding: 10px; background: rgba(255, 235, 59, 0.1); border-radius: 8px;")
        
        mode_layout.addLayout(mode_selection_layout)
        mode_layout.addWidget(self.mode_description)
        mode_group.setLayout(mode_layout)

        # Create tabbed interface for different sections
        self.tab_widget = QTabWidget()
        
        # Tab 1: File Management
        files_tab = QWidget()
        files_layout = QVBoxLayout(files_tab)
        
        # Input Files Section
        input_group = QGroupBox("📁 ไฟล์ที่ต้องใส่ (Input Files)")
        input_layout = QVBoxLayout()
        
        file_controls = QHBoxLayout()
        self.file_btn = QPushButton("เลือกไฟล์ (Select Files)")
        self.file_btn.clicked.connect(self.select_files)
        self.clear_files_btn = QPushButton("ล้างรายการไฟล์ (Clear Files)")
        self.clear_files_btn.setObjectName("clearButton")
        self.clear_files_btn.clicked.connect(self.clear_files)
        file_controls.addWidget(self.file_btn)
        file_controls.addWidget(self.clear_files_btn)
        file_controls.addStretch()
        
        self.files_table = QTableWidget()
        self.files_table.setColumnCount(4)
        self.files_table.setHorizontalHeaderLabels(["ชื่อไฟล์ (File Name)", "ประเภท (Type)", "ขนาด (Size)", "สถานะ (Status)"])
        self.files_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.files_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.files_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.files_table.verticalHeader().setVisible(False)
        
        input_layout.addLayout(file_controls)
        input_layout.addWidget(self.files_table)
        input_group.setLayout(input_layout)
        files_layout.addWidget(input_group)
        
        # Tab 2: Mode Configuration
        config_tab = QWidget()
        config_layout = QVBoxLayout(config_tab)
        
        # Dynamic configuration area
        self.config_group = QGroupBox("⚙️ การตั้งค่าโหมด (Mode Configuration)")
        self.config_layout = QVBoxLayout()
        self.config_group.setLayout(self.config_layout)
        config_layout.addWidget(self.config_group)
        
        # Tab 3: Text Input
        text_tab = QWidget()
        text_layout = QVBoxLayout(text_tab)
        
        text_input_group = QGroupBox("📝 ข้อความที่ต้องการซ่อน (Text to Hide)")
        text_input_layout = QVBoxLayout()
        
        self.text_input = QPlainTextEdit()
        self.text_input.setPlaceholderText("กรอกข้อความที่ต้องการซ่อนที่นี่... (Enter text to hide here...)")
        self.text_input.setMinimumHeight(200)
        
        text_controls = QHBoxLayout()
        load_text_btn = QPushButton("โหลดจากไฟล์ (Load from File)")
        load_text_btn.clicked.connect(self.load_text_from_file)
        save_text_btn = QPushButton("บันทึกข้อความ (Save Text)")
        save_text_btn.clicked.connect(self.save_text_to_file)
        clear_text_btn = QPushButton("ล้างข้อความ (Clear Text)")
        clear_text_btn.setObjectName("clearButton")
        clear_text_btn.clicked.connect(lambda: self.text_input.clear())
        
        text_controls.addWidget(load_text_btn)
        text_controls.addWidget(save_text_btn)
        text_controls.addWidget(clear_text_btn)
        text_controls.addStretch()
        
        text_input_layout.addWidget(self.text_input)
        text_input_layout.addLayout(text_controls)
        text_input_group.setLayout(text_input_layout)
        text_layout.addWidget(text_input_group)
        
        # Add tabs
        self.tab_widget.addTab(files_tab, "📁 Files")
        self.tab_widget.addTab(config_tab, "⚙️ Config")
        self.tab_widget.addTab(text_tab, "📝 Text")
        
        # Workflow Section
        workflow_group = QGroupBox("🔄 ลำดับการทำงาน (Workflow Queue)")
        workflow_layout = QVBoxLayout()
        
        workflow_controls = QHBoxLayout()
        add_to_workflow_btn = QPushButton("เพิ่มเข้าคิว (Add to Queue)")
        add_to_workflow_btn.setObjectName("modeButton")
        add_to_workflow_btn.clicked.connect(self.add_to_workflow)
        
        remove_from_workflow_btn = QPushButton("ลบจากคิว (Remove from Queue)")
        remove_from_workflow_btn.setObjectName("clearButton")
        remove_from_workflow_btn.clicked.connect(self.remove_from_workflow)
        
        clear_workflow_btn = QPushButton("ล้างคิว (Clear Queue)")
        clear_workflow_btn.setObjectName("clearButton")
        clear_workflow_btn.clicked.connect(self.clear_workflow)
        
        workflow_controls.addWidget(add_to_workflow_btn)
        workflow_controls.addWidget(remove_from_workflow_btn)
        workflow_controls.addWidget(clear_workflow_btn)
        workflow_controls.addStretch()
        
        self.workflow_list = QListWidget()
        self.workflow_list.setMinimumHeight(150)
        
        workflow_layout.addLayout(workflow_controls)
        workflow_layout.addWidget(self.workflow_list)
        workflow_group.setLayout(workflow_layout)
        
        # Results Section
        result_group = QGroupBox("📤 ผลลัพธ์และการดำเนินการ (Results & Execution)")
        result_layout = QVBoxLayout()
        
        # Output path selection
        output_path_layout = QHBoxLayout()
        output_path_label = QLabel("ตำแหน่งบันทึกผลลัพธ์ (Output Path):")
        output_path_label.setObjectName("titleLabel")
        self.output_path_display = QLabel("ยังไม่ได้เลือก (Not selected)")
        self.output_path_display.setStyleSheet("color: #cccccc; font-size: 12px;")
        self.select_output_btn = QPushButton("เลือกโฟลเดอร์ (Select Folder)")
        self.select_output_btn.clicked.connect(self.select_output_path)
        
        output_path_layout.addWidget(output_path_label)
        output_path_layout.addWidget(self.output_path_display)
        output_path_layout.addStretch()
        output_path_layout.addWidget(self.select_output_btn)
        
        # Results display
        self.result_display = QTextEdit()
        self.result_display.setReadOnly(True)
        self.result_display.setPlaceholderText("ผลลัพธ์จะแสดงที่นี่ (Results will appear here...)")
        self.result_display.setMinimumHeight(150)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setRange(0, 100)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        
        # Execute controls
        execute_layout = QHBoxLayout()
        self.execute_btn = QPushButton("🚀 เริ่มทำงาน (Execute Workflow)")
        self.execute_btn.setObjectName("executeButton")
        self.execute_btn.clicked.connect(self.execute_workflow)
        
        self.clear_results_btn = QPushButton("ล้างผลลัพธ์ (Clear Results)")
        self.clear_results_btn.setObjectName("clearButton")
        self.clear_results_btn.clicked.connect(lambda: self.result_display.clear())
        
        execute_layout.addWidget(self.execute_btn)
        execute_layout.addStretch()
        execute_layout.addWidget(self.clear_results_btn)
        
        result_layout.addLayout(output_path_layout)
        result_layout.addWidget(self.result_display)
        result_layout.addWidget(self.progress_bar)
        result_layout.addLayout(execute_layout)
        result_group.setLayout(result_layout)
        
        # === สร้าง Extract Mode Section ===
        extract_mode_group = QGroupBox("🔍 เลือกโหมดการถอด (Select Extract Mode)")
        extract_mode_layout = QVBoxLayout()
        
        extract_mode_selection_layout = QHBoxLayout()
        extract_mode_selection_layout.addWidget(QLabel("โหมดการถอด:"))
        self.extract_mode_dropdown = QComboBox()
        self.extract_mode_dropdown.setMinimumWidth(400)
        
        # Add extract modes
        for mode_id, mode_name in self.modes.items():
            extract_name = mode_name.replace("ซ่อน", "ถอด").replace("Hide", "Extract")
            self.extract_mode_dropdown.addItem(f"🔓 {extract_name}", mode_id)
        
        self.extract_mode_dropdown.currentIndexChanged.connect(self.on_extract_mode_changed)
        extract_mode_selection_layout.addWidget(self.extract_mode_dropdown)
        extract_mode_selection_layout.addStretch()
        
        # Extract mode description
        self.extract_mode_description = QLabel()
        self.extract_mode_description.setWordWrap(True)
        self.extract_mode_description.setStyleSheet("color: #4caf50; font-size: 12px; padding: 10px; background: rgba(76, 175, 80, 0.1); border-radius: 8px;")
        
        extract_mode_layout.addLayout(extract_mode_selection_layout)
        extract_mode_layout.addWidget(self.extract_mode_description)
        extract_mode_group.setLayout(extract_mode_layout)
        
        # === สร้าง Tab สำหรับส่วนการถอดข้อมูล ===
        extract_tab_widget = QTabWidget()

        # Tab 1: ไฟล์ที่ต้องการถอด
        extract_files_tab = QWidget()
        extract_files_layout = QVBoxLayout(extract_files_tab)

        # ไฟล์ที่ต้องการถอด
        extract_files_group = QGroupBox("📂 ไฟล์ที่ต้องการถอด (Files to Extract)")
        extract_files_content_layout = QVBoxLayout()

        # ปุ่มเลือกไฟล์
        extract_file_controls = QHBoxLayout()
        self.extract_file_btn = QPushButton("เลือกไฟล์ (Select Files)")
        self.extract_file_btn.clicked.connect(self.select_extract_files)
        self.clear_extract_files_btn = QPushButton("ล้างรายการไฟล์ (Clear Files)")
        self.clear_extract_files_btn.setObjectName("clearButton")
        self.clear_extract_files_btn.clicked.connect(self.clear_extract_files)
        extract_file_controls.addWidget(self.extract_file_btn)
        extract_file_controls.addWidget(self.clear_extract_files_btn)
        extract_file_controls.addStretch()

        # ตารางไฟล์
        self.extract_files_table = QTableWidget()
        self.extract_files_table.setColumnCount(4)
        self.extract_files_table.setHorizontalHeaderLabels(["ชื่อไฟล์ (File Name)", "ประเภท (Type)", "ขนาด (Size)", "สถานะ (Status)"])
        self.extract_files_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.extract_files_table.setSelectionBehavior(QTableWidget.SelectRows)
        self.extract_files_table.setEditTriggers(QTableWidget.NoEditTriggers)
        self.extract_files_table.verticalHeader().setVisible(False)

        # เพิ่มเข้า Group
        extract_files_content_layout.addLayout(extract_file_controls)
        extract_files_content_layout.addWidget(self.extract_files_table)
        extract_files_group.setLayout(extract_files_content_layout)
        extract_files_layout.addWidget(extract_files_group)

        # Tab 2: การตั้งค่าการถอด
        config_extract_tab = QWidget()
        config_extract_layout = QVBoxLayout(config_extract_tab)
        self.extract_config_group = QGroupBox("🔑 การตั้งค่าการถอด (Extract Configuration)")
        self.extract_config_layout = QVBoxLayout()
        self.extract_config_group.setLayout(self.extract_config_layout)
        config_extract_layout.addWidget(self.extract_config_group)

        # Tab 3: ข้อความที่ถอดได้
        extracted_text_tab = QWidget()
        extracted_text_tab_layout = QVBoxLayout(extracted_text_tab)

        extracted_text_group = QGroupBox("📄 ข้อความที่ถอดได้ (Extracted Text)")
        extracted_text_content_layout = QVBoxLayout()

        # กล่องแสดงข้อความ
        self.extracted_text_display = QPlainTextEdit()
        self.extracted_text_display.setPlaceholderText("ข้อความที่ถอดได้จะแสดงที่นี่... (Extracted text will appear here...)")
        self.extracted_text_display.setMinimumHeight(300)
        self.extracted_text_display.setReadOnly(True)

        # ปุ่มควบคุม
        extracted_text_controls = QHBoxLayout()
        copy_text_btn = QPushButton("คัดลอกข้อความ (Copy Text)")
        copy_text_btn.clicked.connect(self.copy_extracted_text)
        save_extracted_btn = QPushButton("บันทึกข้อความ (Save Text)")
        save_extracted_btn.clicked.connect(self.save_extracted_text)
        clear_extracted_btn = QPushButton("ล้างข้อความ (Clear Text)")
        clear_extracted_btn.setObjectName("clearButton")
        clear_extracted_btn.clicked.connect(lambda: self.extracted_text_display.clear())
        extracted_text_controls.addWidget(copy_text_btn)
        extracted_text_controls.addWidget(save_extracted_btn)
        extracted_text_controls.addWidget(clear_extracted_btn)
        extracted_text_controls.addStretch()

        # เพิ่มเข้า Group
        extracted_text_content_layout.addWidget(self.extracted_text_display)
        extracted_text_content_layout.addLayout(extracted_text_controls)
        extracted_text_group.setLayout(extracted_text_content_layout)
        extracted_text_tab_layout.addWidget(extracted_text_group)

        # เพิ่มแท็บทั้งหมด
        extract_tab_widget.addTab(extract_files_tab, "📂 Files")
        extract_tab_widget.addTab(config_extract_tab, "⚙️ Config")
        extract_tab_widget.addTab(extracted_text_tab, "📝 Text")

        # เพิ่มปุ่มเริ่มถอดข้อมูลด้านล่างแท็บ
        extract_execute_layout = QHBoxLayout()
        self.extract_execute_btn = QPushButton("🔓 เริ่มถอดข้อมูล (Start Extraction)")
        self.extract_execute_btn.setObjectName("executeButton")
        self.extract_execute_btn.clicked.connect(self.execute_extraction)
        extract_execute_layout.addWidget(self.extract_execute_btn)
        extract_execute_layout.addStretch()
        
        # สร้าง Horizontal Layout เพื่อจัดวาง extract_mode_group และ extract_tab_widget ข้างกัน
        horizontal_extract_layout = QHBoxLayout()
        horizontal_extract_layout.addWidget(extract_mode_group)
        horizontal_extract_layout.addWidget(extract_tab_widget, 1)
        
        # สร้าง Horizontal Layout เพื่อจัดวาง mode_group และ tab_widget ข้างกัน
        horizontal_layout = QHBoxLayout()
        horizontal_layout.addWidget(mode_group)
        horizontal_layout.addWidget(self.tab_widget, 1)  # 1 = stretch factor

        # เพิ่มเข้าไปใน main layout
        main_layout.addLayout(title_layout)
        main_layout.addLayout(horizontal_layout)
        main_layout.addWidget(workflow_group)
        main_layout.addWidget(result_group)
        main_layout.addLayout(horizontal_extract_layout)
        main_layout.addLayout(extract_execute_layout)
        
        # Set scroll area
        scroll_area.setWidget(container_widget)
        outer_layout = QVBoxLayout(self)
        outer_layout.addWidget(scroll_area)
        self.setLayout(outer_layout)
            
    def copy_extracted_text(self):
        """Copy extracted text to clipboard"""
        from PyQt5.QtWidgets import QApplication
        text = self.extracted_text_display.toPlainText()
        if text.strip():
            clipboard = QApplication.clipboard()
            clipboard.setText(text)
            self.result_display.append("<span style='color: #00ff88;'>✅ คัดลอกข้อความไปยัง clipboard แล้ว</span>")
        else:
            QMessageBox.warning(self, "คำเตือน", "ไม่มีข้อความให้คัดลอก")

    def save_extracted_text(self):
        """Save extracted text to file"""
        text = self.extracted_text_display.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "คำเตือน", "ไม่มีข้อความให้บันทึก")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(self, "บันทึกข้อความที่ถอดได้", "", "Text Files (*.txt);;All Files (*.*)")
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(text)
                    self.result_display.append(f"<span style='color: #00ff88;'>✅ บันทึกข้อความที่ถอดได้ไปยังไฟล์: {os.path.basename(file_path)}</span>")
            except Exception as e:
                QMessageBox.warning(self, "ข้อผิดพลาด", f"ไม่สามารถบันทึกไฟล์ได้: {str(e)}")

    def select_extract_files(self):
        """Select files for extraction"""
        file_types = "All Supported Files (*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.wav *.mp3 *.flac *.ogg *.mp4 *.avi *.mov *.mkv *.docx *.pdf);;Images (*.png *.jpg *.jpeg *.bmp *.gif *.tiff);;Audio (*.wav *.mp3 *.flac *.ogg);;Video (*.mp4 *.avi *.mov *.mkv);;Documents (*.docx *.pdf);;All Files (*.*)"
        files, _ = QFileDialog.getOpenFileNames(self, "เลือกไฟล์สำหรับถอดข้อมูล (Select Files for Extraction)", "", file_types)
        if files:
            self.selected_extract_files.extend(files)
            self.update_extract_files_table()

    def update_extract_files_table(self):
        """Update the extract files table display"""
        self.extract_files_table.setRowCount(len(self.selected_extract_files))
        for row, file_path in enumerate(self.selected_extract_files):
            file_name = os.path.basename(file_path)
            self.extract_files_table.setItem(row, 0, QTableWidgetItem(file_name))
            
            file_ext = os.path.splitext(file_path)[1].lower()
            file_type = "ไม่ระบุ (Unknown)"
            if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']:
                file_type = "ภาพ (Image)"
            elif file_ext in ['.wav', '.mp3', '.flac', '.ogg']:
                file_type = "เสียง (Audio)"
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv']:
                file_type = "วิดีโอ (Video)"
            elif file_ext in ['.docx', '.pdf']:
                file_type = "เอกสาร (Document)"
            self.extract_files_table.setItem(row, 1, QTableWidgetItem(file_type))
            
            try:
                size = os.path.getsize(file_path)
                size_str = self.format_size(size)
            except:
                size_str = "ไม่ทราบ"
            self.extract_files_table.setItem(row, 2, QTableWidgetItem(size_str))
            
            # Status column
            status = "พร้อมถอด (Ready to Extract)"
            self.extract_files_table.setItem(row, 3, QTableWidgetItem(status))

    def clear_extract_files(self):
        """Clear selected extract files"""
        self.selected_extract_files.clear()
        self.extract_files_table.setRowCount(0)
        self.result_display.append("<span style='color: #00d4ff;'>🗑️ รายการไฟล์ถอดข้อมูลถูกล้างแล้ว (Extract file list cleared).</span>")

    def execute_extraction(self):
        """Execute the extraction process"""
        if not self.selected_extract_files:
            QMessageBox.warning(self, "คำเตือน", "กรุณาเลือกไฟล์ที่ต้องการถอดข้อมูล")
            return
        
        # Simulate extraction process
        self.result_display.append("<span style='color: #00d4ff;'>🔓 เริ่มกระบวนการถอดข้อมูล... (Starting extraction process...)</span>")
        
        # Here you would implement the actual extraction logic
        # For now, we'll simulate with dummy text
        extracted_text = f"ตัวอย่างข้อความที่ถอดได้จากไฟล์\nThis is sample extracted text from files\n\nโหมดที่ใช้: {self.modes.get(self.current_extract_mode_id, 'Unknown')}\nจำนวนไฟล์ที่ถอด: {len(self.selected_extract_files)} ไฟล์\nเวลาที่ถอด: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        self.extracted_text_display.setPlainText(extracted_text)
        self.result_display.append("<span style='color: #00ff88;'>✅ ถอดข้อมูลสำเร็จ! ดูผลลัพธ์ในแท็บ 'Text' ของส่วนการถอดข้อมูล</span>")

            # Initialize with first mode
        self.on_mode_changed(0)

        # Initialize extract mode
        self.on_extract_mode_changed(0)

    def on_mode_changed(self, index):
        """Handle mode selection change"""
        mode_id = self.mode_dropdown.itemData(index)
        self.current_mode_id = mode_id
        self.update_mode_description(mode_id)
        self.update_mode_configuration(mode_id)

    def update_mode_description(self, mode_id):
        """Update mode description based on selected mode"""
        descriptions = {
            1: "🔐 เข้ารหัสข้อความด้วย AES → แบ่งครึ่ง → ซ่อนในภาพและเสียง\n✅ ต้องมีไฟล์ภาพและเสียงเพื่อถอดข้อความ",
            2: "📄 สร้างไฟล์ DOCX → ตั้งรหัsผ่าน → เข้ารหัส RSA → ซ่อนใน metadata วิดีโอ\n✅ ต้องมี RSA private key เพื่อถอดรหัส",
            3: "🎛️ เข้ารหัส AES → แบ่ง 3 ส่วน → ซ่อนในภาพ/เสียง/วิดีโอ\n✅ ต้องได้ไฟล์ครบทั้ง 3 ก่อนถึงจะถอดข้อความได้",
            4: "🧬 เข้ารหัส AES → เข้ารหัส key ด้วย RSA → ซ่อนใน metadata\n✅ ต้องมี RSA private key เพื่อถอดรหัส key → ถอดข้อความ",
            5: "🧫 ใช้ GPG เข้ารหัส → ซ่อนใน metadata + EOF embedding\n✅ รองรับการกู้คืนแม้ metadata ถูก strip",
            6: "🧩 เข้ารหัส AES → ซ่อน ciphertext ใน LSB → ซ่อน key ใน metadata → checksum\n✅ บังคับให้ผู้ถอดต้องเข้าถึงทั้งภาพ เสียง และวิดีโอ",
            7: "🔄 base64 + gzip → AES → LSB + metadata + EOF\n✅ มี redundancy และ layer ซ้อนกัน",
            8: "🧠 AES + GPG → กระจาย key และข้อมูลในหลายไฟล์\n✅ กระจาย key และข้อมูลให้ปลอดภัยในหลายไฟล์",
            9: "🌀 ซ่อนข้อความในภาพ → ซ่อนภาพในเสียง → ซ่อนเสียงในวิดีโอ\n✅ เหมือน 'ตุ๊กตารัสเซีย' ซ้อนกันหลายชั้น",
            10: "🧾 โหมด 10: Split + Layered + Time-lock"
        }
        
        self.mode_description.setText(descriptions.get(mode_id, "ไม่พบคำอธิบาย"))

    def update_mode_configuration(self, mode_id):
        """Update configuration UI based on selected mode"""
        # Clear existing configuration completely
        self.clear_config_layout()
        
        # Create configuration based on mode
        if mode_id == 1:
            self.create_mode1_config()
        elif mode_id == 2:
            self.create_mode2_config()
        elif mode_id == 3:
            self.create_mode3_config()
        elif mode_id == 4:
            self.create_mode4_config()
        elif mode_id == 5:
            self.create_mode5_config()
        elif mode_id == 6:
            self.create_mode6_config()
        elif mode_id == 7:
            self.create_mode7_config()
        elif mode_id == 8:
            self.create_mode8_config()
        elif mode_id == 9:
            self.create_mode9_config()
        elif mode_id == 10:
            self.create_mode10_config()

    def clear_config_layout(self):
        """Completely clear the configuration layout"""
        # Remove all widgets and layouts recursively
        def clear_layout(layout):
            if layout is not None:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget is not None:
                        widget.deleteLater()
                    else:
                        # If it's a layout, clear it recursively
                        child_layout = item.layout()
                        if child_layout is not None:
                            clear_layout(child_layout)
        
        clear_layout(self.config_layout)
        
        # Force update the UI
        self.config_group.update()

    def create_mode1_config(self):
        """Configuration for Mode 1: AES + Split Half"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.aes_password = QLineEdit()
        self.aes_password.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password, 1, 1)
        
        self.random_aes = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes.toggled.connect(lambda checked: self.aes_password.setEnabled(not checked))
        layout.addWidget(self.random_aes, 2, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 3, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (สำหรับซ่อนครึ่งแรก)\n• 1 ไฟล์เสียง (สำหรับซ่อนครึ่งหลัง)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 4, 0, 1, 2)
        
        self.config_layout.addLayout(layout)



    def create_mode2_config(self):
        """Configuration for Mode 2: DOCX + RSA + Video Metadata"""
        layout = QGridLayout()
        
        # DOCX Configuration
        layout.addWidget(QLabel("📄 การตั้งค่า DOCX:"), 0, 0, 1, 2)
        layout.addWidget(QLabel("รหัสผ่าน DOCX:"), 1, 0)
        self.docx_password = QLineEdit()
        self.docx_password.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.docx_password, 1, 1)
        
        self.random_docx = QCheckBox("สุ่มรหัสผ่าน DOCX")
        self.random_docx.toggled.connect(lambda checked: self.docx_password.setEnabled(not checked))
        layout.addWidget(self.random_docx, 2, 0, 1, 2)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA:"), 3, 0, 1, 2)
        layout.addWidget(QLabel("RSA Public Key:"), 4, 0)
        
        rsa_layout = QHBoxLayout()
        self.rsa_public_key = QLineEdit()
        self.rsa_public_key.setPlaceholderText("เส้นทางไฟล์ public key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.rsa_public_key, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.rsa_public_key)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 4, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์วิดีโอ (สำหรับซ่อนใน metadata)\n• RSA public key file")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.config_layout.addLayout(layout)


    def create_mode3_config(self):
        """Configuration for Mode 3: AES + Split 3 Parts"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.aes_password_m3 = QLineEdit()
        self.aes_password_m3.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password_m3, 1, 1)
        
        self.random_aes_m3 = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes_m3.toggled.connect(lambda checked: self.aes_password_m3.setEnabled(not checked))
        layout.addWidget(self.random_aes_m3, 2, 0, 1, 2)
        
        # Split Configuration
        layout.addWidget(QLabel("✂️ การแบ่งข้อความ:"), 3, 0, 1, 2)
        
        self.equal_split = QRadioButton("แบ่งเท่าๆ กัน 3 ส่วน")
        self.equal_split.setChecked(True)
        layout.addWidget(self.equal_split, 4, 0, 1, 2)
        
        self.custom_split = QRadioButton("กำหนดสัดส่วนเอง")
        layout.addWidget(self.custom_split, 5, 0, 1, 2)
        
        # Custom split ratios
        split_ratio_layout = QHBoxLayout()
        split_ratio_layout.addWidget(QLabel("ภาพ:"))
        self.image_ratio = QSpinBox()
        self.image_ratio.setRange(1, 100)
        self.image_ratio.setValue(33)
        self.image_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.image_ratio)
        
        split_ratio_layout.addWidget(QLabel("เสียง:"))
        self.audio_ratio = QSpinBox()
        self.audio_ratio.setRange(1, 100)
        self.audio_ratio.setValue(33)
        self.audio_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.audio_ratio)
        
        split_ratio_layout.addWidget(QLabel("วิดีโอ:"))
        self.video_ratio = QSpinBox()
        self.video_ratio.setRange(1, 100)
        self.video_ratio.setValue(34)
        self.video_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.video_ratio)
        
        self.custom_split.toggled.connect(lambda checked: [
            self.image_ratio.setEnabled(checked),
            self.audio_ratio.setEnabled(checked),
            self.video_ratio.setEnabled(checked)
        ])
        
        layout.addLayout(split_ratio_layout, 6, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 7, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ\n• 1 ไฟล์เสียง\n• 1 ไฟล์วิดีโอ")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 8, 0, 1, 2)
        
        self.config_layout.addLayout(layout)





    def create_mode4_config(self):
        """Configuration for Mode 4: AES + RSA + Metadata"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        self.auto_aes_key = QCheckBox("สุ่ม AES key อัตโนมัติ")
        self.auto_aes_key.setChecked(True)
        layout.addWidget(self.auto_aes_key, 1, 0, 1, 2)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA:"), 2, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Public Key:"), 3, 0)
        rsa_layout = QHBoxLayout()
        self.rsa_public_key_m4 = QLineEdit()
        self.rsa_public_key_m4.setPlaceholderText("เส้นทางไฟล์ public key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.rsa_public_key_m4, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.rsa_public_key_m4)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 3, 1)
        
        # Metadata Options
        layout.addWidget(QLabel("📋 ตัวเลือก Metadata:"), 4, 0, 1, 2)
        
        self.preserve_original = QCheckBox("เก็บ metadata เดิมไว้")
        self.preserve_original.setChecked(True)
        layout.addWidget(self.preserve_original, 5, 0, 1, 2)
        
        self.add_dummy_data = QCheckBox("เพิ่มข้อมูลปลอมเพื่อปกปิด")
        layout.addWidget(self.add_dummy_data, 6, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 7, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์เสียง (สำหรับ ciphertext)\n• 1 ไฟล์วิดีโอ (สำหรับ RSA-encrypted key)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 8, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode5_config(self):
        """Configuration for Mode 5: GPG + Metadata + EOF"""
        layout = QGridLayout()
        
        # GPG Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า GPG:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("GPG Public Key:"), 1, 0)
        gpg_layout = QHBoxLayout()
        self.gpg_public_key = QLineEdit()
        self.gpg_public_key.setPlaceholderText("เส้นทางไฟล์ GPG public key")
        gpg_browse_btn = QPushButton("เลือกไฟล์")
        gpg_browse_btn.clicked.connect(lambda: self.browse_file(self.gpg_public_key, "GPG Key Files (*.asc *.gpg)"))
        gpg_layout.addWidget(self.gpg_public_key)
        gpg_layout.addWidget(gpg_browse_btn)
        layout.addLayout(gpg_layout, 1, 1)
        
        # EOF Embedding Options
        layout.addWidget(QLabel("➕ ตัวเลือก EOF Embedding:"), 2, 0, 1, 2)
        
        self.eof_format = QComboBox()
        self.eof_format.addItems(["ZIP → MP4", "RAR → AVI", "7Z → MKV", "TAR → MOV"])
        layout.addWidget(QLabel("รูปแบบ EOF:"), 3, 0)
        layout.addWidget(self.eof_format, 3, 1)
        
        self.redundancy_level = QSpinBox()
        self.redundancy_level.setRange(1, 5)
        self.redundancy_level.setValue(2)
        layout.addWidget(QLabel("ระดับ Redundancy:"), 4, 0)
        layout.addWidget(self.redundancy_level, 4, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (metadata)\n• 1 ไฟล์เสียง (metadata)\n• 1 ไฟล์วิดีโอ (EOF embedding)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode6_config(self):
        """Configuration for Mode 6: AES + LSB + Metadata + Checksum"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.aes_password_m6 = QLineEdit()
        self.aes_password_m6.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password_m6, 1, 1)
        
        self.random_aes_m6 = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes_m6.toggled.connect(lambda checked: self.aes_password_m6.setEnabled(not checked))
        layout.addWidget(self.random_aes_m6, 2, 0, 1, 2)
        
        # LSB Configuration
        layout.addWidget(QLabel("🖼️ การตั้งค่า LSB:"), 3, 0, 1, 2)
        
        self.lsb_bits = QSpinBox()
        self.lsb_bits.setRange(1, 4)
        self.lsb_bits.setValue(2)
        layout.addWidget(QLabel("จำนวน LSB bits:"), 4, 0)
        layout.addWidget(self.lsb_bits, 4, 1)
        
        self.lsb_channels = QComboBox()
        self.lsb_channels.addItems(["RGB ทั้งหมด", "เฉพาะ Red", "เฉพาะ Green", "เฉพาะ Blue"])
        layout.addWidget(QLabel("ช่องสี:"), 5, 0)
        layout.addWidget(self.lsb_channels, 5, 1)
        
        # Checksum Configuration
        layout.addWidget(QLabel("🧪 การตั้งค่า Checksum:"), 6, 0, 1, 2)
        
        self.checksum_algorithm = QComboBox()
        self.checksum_algorithm.addItems(["SHA-256", "SHA-512", "MD5", "CRC32"])
        layout.addWidget(QLabel("อัลกอริทึม:"), 7, 0)
        layout.addWidget(self.checksum_algorithm, 7, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 8, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (LSB)\n• 1 ไฟล์เสียง (AES key metadata)\n• 1 ไฟล์วิดีโอ (checksum metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 9, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode7_config(self):
        """Configuration for Mode 7: Multi-layer Transformation"""
        layout = QGridLayout()
        
        # Transformation Options
        layout.addWidget(QLabel("🧪 ตัวเลือกการแปลง:"), 0, 0, 1, 2)
        
        self.use_base64 = QCheckBox("ใช้ Base64 encoding")
        self.use_base64.setChecked(True)
        layout.addWidget(self.use_base64, 1, 0, 1, 2)
        
        self.use_gzip = QCheckBox("ใช้ Gzip compression")
        self.use_gzip.setChecked(True)
        layout.addWidget(self.use_gzip, 2, 0, 1, 2)
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 3, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 4, 0)
        self.aes_password_m7 = QLineEdit()
        self.aes_password_m7.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password_m7, 4, 1)
        
        self.random_aes_m7 = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes_m7.toggled.connect(lambda checked: self.aes_password_m7.setEnabled(not checked))
        layout.addWidget(self.random_aes_m7, 5, 0, 1, 2)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA (สำหรับ key backup):"), 6, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Public Key:"), 7, 0)
        rsa_layout = QHBoxLayout()
        self.rsa_public_key_m7 = QLineEdit()
        self.rsa_public_key_m7.setPlaceholderText("เส้นทางไฟล์ public key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.rsa_public_key_m7, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.rsa_public_key_m7)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 7, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 8, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (LSB)\n• 1 ไฟล์วิดีโอ (metadata)\n• 1 ไฟล์เสียง (EOF backup)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 9, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode8_config(self):
        """Configuration for Mode 8: AES + GPG + Multi Media"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.aes_password_m8 = QLineEdit()
        self.aes_password_m8.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password_m8, 1, 1)
        
        self.random_aes_m8 = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes_m8.toggled.connect(lambda checked: self.aes_password_m8.setEnabled(not checked))
        layout.addWidget(self.random_aes_m8, 2, 0, 1, 2)
        
        # GPG Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า GPG:"), 3, 0, 1, 2)
        
        layout.addWidget(QLabel("GPG Public Key:"), 4, 0)
        gpg_layout = QHBoxLayout()
        self.gpg_public_key_m8 = QLineEdit()
        self.gpg_public_key_m8.setPlaceholderText("เส้นทางไฟล์ GPG public key")
        gpg_browse_btn = QPushButton("เลือกไฟล์")
        gpg_browse_btn.clicked.connect(lambda: self.browse_file(self.gpg_public_key_m8, "GPG Key Files (*.asc *.gpg)"))
        gpg_layout.addWidget(self.gpg_public_key_m8)
        gpg_layout.addWidget(gpg_browse_btn)
        layout.addLayout(gpg_layout, 4, 1)
        
        # Distribution Strategy
        layout.addWidget(QLabel("📊 กลยุทธ์การกระจายข้อมูล:"), 5, 0, 1, 2)
        
        self.distribution_strategy = QComboBox()
        self.distribution_strategy.addItems([
            "แบ่งเท่าๆ กัน",
            "ตามขนาดไฟล์",
            "ตามความสำคัญ",
            "สุ่มการกระจาย"
        ])
        layout.addWidget(QLabel("กลยุทธ์:"), 6, 0)
        layout.addWidget(self.distribution_strategy, 6, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 7, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์วิดีโอ (ciphertext metadata)\n• 1 ไฟล์ภาพ (GPG-encrypted key)\n• 1 ไฟล์เสียง (hash EOF)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 8, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode9_config(self):
        """Configuration for Mode 9: Nested Stego"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.aes_password_m9 = QLineEdit()
        self.aes_password_m9.setPlaceholderText("กรอกรหัสผ่าน หรือ ปล่อยว่างเพื่อสุ่ม")
        layout.addWidget(self.aes_password_m9, 1, 1)
        
        self.random_aes_m9 = QCheckBox("สุ่มรหัสผ่าน AES")
        self.random_aes_m9.toggled.connect(lambda checked: self.aes_password_m9.setEnabled(not checked))
        layout.addWidget(self.random_aes_m9, 2, 0, 1, 2)
        
        # Nesting Configuration
        layout.addWidget(QLabel("🌀 การตั้งค่าการซ้อน:"), 3, 0, 1, 2)
        
        self.nesting_levels = QSpinBox()
        self.nesting_levels.setRange(2, 5)
        self.nesting_levels.setValue(3)
        layout.addWidget(QLabel("จำนวนชั้นการซ้อน:"), 4, 0)
        layout.addWidget(self.nesting_levels, 4, 1)
        
        self.encoding_method = QComboBox()
        self.encoding_method.addItems(["Base64", "Hex", "Binary", "Custom"])
        layout.addWidget(QLabel("วิธีการ encoding:"), 5, 0)
        layout.addWidget(self.encoding_method, 5, 1)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA:"), 6, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Public Key:"), 7, 0)
        rsa_layout = QHBoxLayout()
        self.rsa_public_key_m9 = QLineEdit()
        self.rsa_public_key_m9.setPlaceholderText("เส้นทางไฟล์ public key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.rsa_public_key_m9, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.rsa_public_key_m9)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 7, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 8, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (ชั้นใน)\n• 1 ไฟล์เสียง (ชั้นกลาง)\n• 1 ไฟล์วิดีโอ (ชั้นนอก + key metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 9, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def create_mode10_config(self):
        """Configuration for Mode 10: Split + Layered + Time-lock"""
        layout = QGridLayout()
        
        # Split Configuration
        layout.addWidget(QLabel("✂️ การตั้งค่าการแบ่ง:"), 0, 0, 1, 2)
        
        self.split_parts = QSpinBox()
        self.split_parts.setRange(2, 10)
        self.split_parts.setValue(3)
        layout.addWidget(QLabel("จำนวนส่วนที่แบ่ง:"), 1, 0)
        layout.addWidget(self.split_parts, 1, 1)
        
        # AES Configuration for each part
        layout.addWidget(QLabel("🔐 การตั้งค่า AES (หลาย keys):"), 2, 0, 1, 2)
        
        self.use_different_keys = QCheckBox("ใช้ AES key ต่างกันสำหรับแต่ละส่วน")
        self.use_different_keys.setChecked(True)
        layout.addWidget(self.use_different_keys, 3, 0, 1, 2)
        
        layout.addWidget(QLabel("Master Password:"), 4, 0)
        self.master_password = QLineEdit()
        self.master_password.setPlaceholderText("รหัสผ่านหลักสำหรับสร้าง keys")
        layout.addWidget(self.master_password, 4, 1)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA (หลาย keys):"), 5, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Public Keys:"), 6, 0)
        rsa_keys_layout = QVBoxLayout()
        
        self.rsa_key_list = QListWidget()
        self.rsa_key_list.setMaximumHeight(100)
        rsa_keys_layout.addWidget(self.rsa_key_list)
        
        rsa_buttons_layout = QHBoxLayout()
        add_rsa_btn = QPushButton("เพิ่ม RSA Key")
        add_rsa_btn.clicked.connect(self.add_rsa_key)
        remove_rsa_btn = QPushButton("ลบ RSA Key")
        remove_rsa_btn.setObjectName("clearButton")
        remove_rsa_btn.clicked.connect(self.remove_rsa_key)
        rsa_buttons_layout.addWidget(add_rsa_btn)
        rsa_buttons_layout.addWidget(remove_rsa_btn)
        rsa_keys_layout.addLayout(rsa_buttons_layout)
        
        layout.addLayout(rsa_keys_layout, 6, 1)
        
        # Time-lock Configuration
        layout.addWidget(QLabel("⏳ การตั้งค่า Time-lock:"), 7, 0, 1, 2)
        
        self.enable_timelock = QCheckBox("เปิดใช้งาน Time-lock")
        layout.addWidget(self.enable_timelock, 8, 0, 1, 2)
        
        timelock_layout = QHBoxLayout()
        timelock_layout.addWidget(QLabel("ปลดล็อกหลังจาก:"))
        self.timelock_hours = QSpinBox()
        self.timelock_hours.setRange(0, 8760)  # Up to 1 year
        self.timelock_hours.setValue(24)
        self.timelock_hours.setEnabled(False)
        timelock_layout.addWidget(self.timelock_hours)
        timelock_layout.addWidget(QLabel("ชั่วโมง"))
        
        self.enable_timelock.toggled.connect(lambda checked: self.timelock_hours.setEnabled(checked))
        layout.addLayout(timelock_layout, 9, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 10, 0, 1, 2)
        req_label = QLabel("• หลายไฟล์ภาพ/เสียง/วิดีโอ (ตามจำนวนส่วนที่แบ่ง)\n• หลาย RSA public keys\n• 1 ไฟล์ ZIP (สำหรับ metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 11, 0, 1, 2)
        
        self.config_layout.addLayout(layout)

    def browse_file(self, line_edit, file_filter):
        """Browse and select a file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "เลือกไฟล์", "", file_filter)
        if file_path:
            line_edit.setText(file_path)

    def browse_rsa_public_key(self):
        """Browse RSA public key file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "เลือก RSA Public Key", "", "RSA Key Files (*.pem *.key *.pub);;All Files (*.*)")
        if file_path:
            self.rsa_public_key.setText(file_path)

    def add_rsa_key(self):
        """Add RSA key to the list"""
        file_path, _ = QFileDialog.getOpenFileName(self, "เลือก RSA Public Key", "", "RSA Key Files (*.pem *.key *.pub);;All Files (*.*)")
        if file_path:
            self.rsa_key_list.addItem(os.path.basename(file_path) + " → " + file_path)

    def remove_rsa_key(self):
        """Remove selected RSA key from the list"""
        current_row = self.rsa_key_list.currentRow()
        if current_row >= 0:
            self.rsa_key_list.takeItem(current_row)

    def select_files(self):
        """Select input files"""
        file_types = "All Supported Files (*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.wav *.mp3 *.flac *.ogg *.mp4 *.avi *.mov *.mkv *.docx *.pdf);;Images (*.png *.jpg *.jpeg *.bmp *.gif *.tiff);;Audio (*.wav *.mp3 *.flac *.ogg);;Video (*.mp4 *.avi *.mov *.mkv);;Documents (*.docx *.pdf);;All Files (*.*)"
        files, _ = QFileDialog.getOpenFileNames(self, "เลือกไฟล์ (Select Files)", "", file_types)
        if files:
            self.selected_files.extend(files)
            self.update_files_table()

    def update_files_table(self):
        """Update the files table display"""
        self.files_table.setRowCount(len(self.selected_files))
        for row, file_path in enumerate(self.selected_files):
            file_name = os.path.basename(file_path)
            self.files_table.setItem(row, 0, QTableWidgetItem(file_name))
            
            file_ext = os.path.splitext(file_path)[1].lower()
            file_type = "ไม่ระบุ (Unknown)"
            if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']:
                file_type = "ภาพ (Image)"
            elif file_ext in ['.wav', '.mp3', '.flac', '.ogg']:
                file_type = "เสียง (Audio)"
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv']:
                file_type = "วิดีโอ (Video)"
            elif file_ext in ['.docx', '.pdf']:
                file_type = "เอกสาร (Document)"
            self.files_table.setItem(row, 1, QTableWidgetItem(file_type))
            
            try:
                size = os.path.getsize(file_path)
                size_str = self.format_size(size)
            except:
                size_str = "ไม่ทราบ"
            self.files_table.setItem(row, 2, QTableWidgetItem(size_str))
            
            # Status column
            status = "พร้อม (Ready)"
            self.files_table.setItem(row, 3, QTableWidgetItem(status))

    def format_size(self, size):
        """Format file size in human readable format"""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} PB"

    def clear_files(self):
        """Clear selected files"""
        self.selected_files.clear()
        self.files_table.setRowCount(0)
        self.result_display.append("<span style='color: #00d4ff;'>🗑️ รายการไฟล์ถูกล้างแล้ว (File list cleared).</span>")

    def load_text_from_file(self):
        """Load text from file"""
        file_path, _ = QFileDialog.getOpenFileName(self, "โหลดข้อความจากไฟล์", "", "Text Files (*.txt);;All Files (*.*)")
        if file_path:
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    self.text_input.setPlainText(content)
                    self.result_display.append(f"<span style='color: #00ff88;'>✅ โหลดข้อความจากไฟล์: {os.path.basename(file_path)}</span>")
            except Exception as e:
                QMessageBox.warning(self, "ข้อผิดพลาด", f"ไม่สามารถโหลดไฟล์ได้: {str(e)}")

    def save_text_to_file(self):
        """Save text to file"""
        if not self.text_input.toPlainText().strip():
            QMessageBox.warning(self, "คำเตือน", "ไม่มีข้อความให้บันทึก")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(self, "บันทึกข้อความ", "", "Text Files (*.txt);;All Files (*.*)")
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(self.text_input.toPlainText())
                    self.result_display.append(f"<span style='color: #00ff88;'>✅ บันทึกข้อความไปยังไฟล์: {os.path.basename(file_path)}</span>")
            except Exception as e:
                QMessageBox.warning(self, "ข้อผิดพลาด", f"ไม่สามารถบันทึกไฟล์ได้: {str(e)}")

    def select_output_path(self):
        """Select output directory"""
        folder = QFileDialog.getExistingDirectory(self, "เลือกตำแหน่งบันทึกไฟล์ผลลัพธ์ (Select Output Folder)")
        if folder:
            self.output_dir = folder
            self.output_path = folder
            self.output_path_display.setText(os.path.basename(folder))
            self.result_display.append(f"<span style='color: #00ff88;'>✅ ตำแหน่งบันทึกผลลัพธ์: {folder}</span>")
        else:
            self.output_path = ""
            self.output_path_display.setText("ยังไม่ได้เลือก (Not selected)")

    def add_to_workflow(self):
        """Add current configuration to workflow"""
        if not self.text_input.toPlainText().strip():
            QMessageBox.warning(self, "คำเตือน", "กรุณากรอกข้อความที่ต้องการซ่อน")
            return
            
        if not self.selected_files:
            QMessageBox.warning(self, "คำเตือน", "กรุณาเลือกไฟล์ที่ต้องการใช้")
            return
            
        # Collect current configuration
        config = self.collect_current_config()
        
        # Create workflow item
        mode_name = self.modes[self.current_mode_id]
        workflow_item = AdvancedSteganoWorkflowItem(
            self.current_mode_id,
            mode_name,
            config
        )
        workflow_item.source_files = self.selected_files.copy()
        
        self.workflow_items.append(workflow_item)
        
        # Add to display
        display_text = f"[{len(self.workflow_items)}] {mode_name}"
        if config.get('text_length'):
            display_text += f" ({config['text_length']} chars)"
        
        self.workflow_list.addItem(display_text)
        self.result_display.append(f"<span style='color: #00d4ff;'>➕ เพิ่มเข้าคิว: {mode_name}</span>")

    def collect_current_config(self):
        """Collect current configuration based on selected mode"""
        config = {
            'mode_id': self.current_mode_id,
            'text': self.text_input.toPlainText(),
            'text_length': len(self.text_input.toPlainText()),
            'files': self.selected_files.copy(),
            'timestamp': datetime.now().isoformat()
        }
        
        # Add mode-specific configuration
        if self.current_mode_id == 1:
            config.update({
                'aes_password': self.aes_password.text() if hasattr(self, 'aes_password') else '',
                'random_aes': self.random_aes.isChecked() if hasattr(self, 'random_aes') else False
            })
        elif self.current_mode_id == 2:
            config.update({
                'docx_password': self.docx_password.text() if hasattr(self, 'docx_password') else '',
                'random_docx': self.random_docx.isChecked() if hasattr(self, 'random_docx') else False,
                'rsa_public_key': self.rsa_public_key.text() if hasattr(self, 'rsa_public_key') else ''
            })
        # Add more mode configurations as needed...
        
        return config

    def remove_from_workflow(self):
        """Remove selected item from workflow"""
        current_row = self.workflow_list.currentRow()
        if current_row >= 0:
            item_text = self.workflow_list.item(current_row).text()
            self.workflow_list.takeItem(current_row)
            self.workflow_items.pop(current_row)
            self.result_display.append(f"<span style='color: #ff4444;'>➖ ลบจากคิว: {item_text}</span>")
        else:
            QMessageBox.warning(self, "คำเตือน", "กรุณาเลือกรายการที่จะลบ")

    def clear_workflow(self):
        """Clear all workflow items"""
        self.workflow_list.clear()
        self.workflow_items.clear()
        self.result_display.append("<span style='color: #00d4ff;'>🗑️ ล้างคิวทั้งหมดแล้ว (Workflow queue cleared).</span>")

    def execute_workflow(self):
        """Execute the workflow with real steganography operations"""
        if not self.workflow_items:
            QMessageBox.warning(self, "คำเตือน", "ไม่มีรายการในคิวการทำงาน")
            return
            
        if not self.output_path:
            QMessageBox.warning(self, "คำเตือน", "กรุณาเลือกตำแหน่งบันทึกผลลัพธ์")
            return

        self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.execute_btn.setEnabled(False)
        
        self.result_display.append("<span style='color: #00d4ff;'>🚀 เริ่มดำเนินการตามคิว... (Starting workflow execution...)</span>")
        
        total_items = len(self.workflow_items)
        for i, item in enumerate(self.workflow_items):
            try:
                progress_value = int(((i + 1) / total_items) * 100)
                self.progress_bar.setValue(progress_value)
                self.result_display.append(f"<span style='color: #ffeb3b;'>⚙️ กำลังดำเนินการ [{i+1}/{total_items}]: {item.mode_name}</span>")
                
                # Execute mode-specific steganography
                if item.mode_id == 1:
                    self.execute_mode1(item)
                elif item.mode_id == 2:
                    self.execute_mode2(item)
                elif item.mode_id == 3:
                    self.execute_mode3(item)
                elif item.mode_id == 4:
                    self.execute_mode4(item)
                elif item.mode_id == 5:
                    self.execute_mode5(item)
                elif item.mode_id == 6:
                    self.execute_mode6(item)
                elif item.mode_id == 7:
                    self.execute_mode7(item)
                elif item.mode_id == 8:
                    self.execute_mode8(item)
                elif item.mode_id == 9:
                    self.execute_mode9(item)
                elif item.mode_id == 10:
                    self.execute_mode10(item)
                else:
                    raise ValueError(f"Unknown mode ID: {item.mode_id}")
                
                self.result_display.append(f"<span style='color: #00ff88;'>✅ สำเร็จ: {item.mode_name}</span>")
                
            except Exception as e:
                self.result_display.append(f"<span style='color: #ff4444;'>❌ ข้อผิดพลาดใน {item.mode_name}: {str(e)}</span>")
                break  # Stop processing on error
            
            # Process UI events to keep the interface responsive
            QTimer.singleShot(100, lambda: None)
        
        self.result_display.append("<span style='color: #00ff88;'>🎉 ดำเนินการเสร็จสมบูรณ์! (All operations completed successfully!)</span>")
        self.progress_bar.setVisible(False)
        self.execute_btn.setEnabled(True)






    def execute_mode1(self, item):
        """Mode 1: Encrypt with AES → Split into 2 parts → Hide in Image (LSB) + Audio (LSB)"""
        config = item.config
        text = config.get('text', '').strip()
        files = item.source_files
        output_dir = self.output_path

        if not text:
            raise ValueError("ไม่พบข้อความสำหรับซ่อน")

        if len(files) < 2:
            raise ValueError("ต้องเลือกอย่างน้อย 1 ไฟล์ภาพ และ 1 ไฟล์เสียง")

        # กรองไฟล์ตามประเภท
        image_file = None
        audio_file = None
        for f in files:
            ext = os.path.splitext(f)[1].lower()
            if ext in ['.png', '.jpg', '.jpeg'] and not image_file:
                image_file = f
            elif ext in ['.wav', '.mp3', '.flac', '.ogg', '.aac'] and not audio_file:
                audio_file = f

        if not image_file:
            raise ValueError("ไม่พบไฟล์ภาพ (รองรับ .png, .jpg)")
        if not audio_file:
            raise ValueError("ไม่พบไฟล์เสียง (รองรับ .wav, .mp3)")

        try:
            # --- ขั้นตอนเดียวกับ Stego.hide() ---
            print("[INFO] Encrypting and hiding data...")

            # กำหนด key
            if config.get('random_aes', False):
                key_str = self.gen_secure_key()
            else:
                password = config.get('aes_password', '').strip()
                if not password:
                    raise ValueError("กรุณากรอกรหัสผ่าน AES หรือเลือก 'สุ่ม'")
                key_str = self.stretch_key(password)  # ยืดให้ได้ 32 ตัว

            # เข้ารหัส
            iv_b64, ct_b64, _, encrypted_b64 = self.encrypt_aes(text, key_str)

            key_b64 = base64.b64encode(key_str.encode()).decode()
            self.result_display.append(f"<span style='color: #00d4ff;'>🔑 Key (string): {key_str}</span>")
            # self.result_display.append(f"<span style='color: #00d4ff;'>🔒 Encrypted (Base64): {encrypted_b64[:60]}...</span>")

            # แบ่งข้อความ
            p1, p2 = self.split_msg(encrypted_b64, 2)

            # ตั้งชื่อไฟล์ผลลัพธ์
            out_img_name = f"stego_img_{uuid.uuid4().hex[:8]}.png"
            out_img_path = os.path.join(output_dir, out_img_name)
            
            # out_audio_dir = os.path.join(output_dir, "audio_output")
            # os.makedirs(out_audio_dir, exist_ok=True)

            # ซ่อนในภาพ
            success_img = self.hide_lsb_image(image_file, p1, out_img_path)
            if not success_img:
                raise RuntimeError("ล้มเหลวในการซ่อนข้อมูลในภาพ")

            # ซ่อนในเสียง
            # out_audio_path = self.hide_lsb_audio(audio_file, p2, out_audio_dir)
            out_audio_path = self.hide_lsb_audio(audio_file, p2, output_dir)

            if not out_audio_path:
                raise RuntimeError("ล้มเหลวในการซ่อนข้อมูลในเสียง")

           
            
           

            self.result_display.append(f"<span style='color: #00ff88;'>✅ ซ่อนข้อมูลสำเร็จ!</span>")
            self.result_display.append(f"<span style='color: #00d4ff;'>🖼️ ภาพ: {out_img_name}</span>")
            self.result_display.append(f"<span style='color: #00d4ff;'>🎵 เสียง: {os.path.basename(out_audio_path)}</span>")
           

        except Exception as e:
            self.result_display.append(f"<span style='color: #ff4444;'>❌ ข้อผิดพลาด: {str(e)}</span>")
            raise




    def execute_mode2(self, item):
        """Mode 2: DOCX + RSA + Video Metadata"""
        try:
            if not self.output_dir:
                raise ValueError("กรุณาเลือกตำแหน่งบันทึกผลลัพธ์ก่อนดำเนินการ")
        
            self.result_display.append("🚀 เริ่มกระบวนการเข้ารหัสและซ่อนข้อมูล...")
            
            # ดึงค่าจาก UI
            docx_password = self.docx_password.text() if not self.random_docx.isChecked() else None
            rsa_key_path = self.rsa_public_key.text()
            video_file = self.selected_files[0] if self.selected_files else None
            
            # ตรวจสอบไฟล์วิดีโอ
            if not video_file or not os.path.exists(video_file):
                raise FileNotFoundError("ไม่พบไฟล์วิดีโอที่เลือก")
            
            # สร้างไฟล์ DOCX
            doc = Document()
            doc.add_paragraph(self.text_input.toPlainText())
            docx_path = os.path.join(self.output_dir, "secret.docx")
            doc.save(docx_path)
            self.result_display.append(f"✅ สร้างไฟล์ DOCX ชั่วคราว: {docx_path}")
            
            # ตั้งรหัสผ่าน DOCX
            if not docx_password:
                docx_password = self.generate_secure_password(16)
            protected_docx = os.path.join(self.output_dir, "protected.docx")
            
            with open(docx_path, "rb") as f_in:
                with open(protected_docx, "wb") as f_out:
                    office_file = msoffcrypto.OfficeFile(f_in)
                    office_file.encrypt(docx_password, f_out)
            self.result_display.append(f"🔒 ตั้งรหัสผ่าน DOCX เรียบร้อย: {protected_docx}")
            
            # โหลด RSA public key
            if not os.path.exists(rsa_key_path):
                raise FileNotFoundError(f"ไม่พบไฟล์ RSA public key: {rsa_key_path}")
            
            with open(rsa_key_path, "rb") as key_file:
                public_key = serialization.load_pem_public_key(key_file.read())
            
            # เข้ารหัสรหัสผ่านด้วย RSA
            encrypted_password = public_key.encrypt(
                docx_password.encode(),
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            encoded_encrypted_password = base64.b64encode(encrypted_password).decode()
            self.result_display.append(f"🔐 รหัสผ่านถูกเข้ารหัสด้วย RSA แล้ว (Base64): {encoded_encrypted_password}")
            
            # ฝังข้อมูลใน metadata วิดีโอ
            output_video = os.path.join(self.output_dir, f"stego_{os.path.basename(video_file)}")
            
            result = subprocess.run([
                "ffmpeg", "-i", video_file,
                "-metadata", f"comment={encoded_encrypted_password}",
                "-codec", "copy", "-y", output_video
            ], capture_output=True, text=True, check=True)
            
            self.result_display.append(f"🎥 ฝังข้อมูลใน metadata วิดีโอสำเร็จ: {output_video}")
            
            # ล้างไฟล์ชั่วคราว
            if os.path.exists(docx_path):
                os.remove(docx_path)
                
        except Exception as e:
            self.result_display.append(f"❌ ข้อผิดพลาด: {str(e)}")
            self.progress_bar.setValue(0)
        else:
            self.progress_bar.setValue(100)
            self.result_display.append("✅ กระบวนการสำเร็จสมบูรณ์!")











    def execute_mode3(self, item):
        """Mode 3: Encrypt with AES → Split into 3 parts → Hide in Image + Audio + Video"""
        config = item.config
        text = config.get('text', '').strip()
        files = item.source_files
        output_dir = self.output_path

        if not text:
            raise ValueError("ไม่พบข้อความสำหรับซ่อน")

        if len(files) < 3:
            raise ValueError("ต้องเลือกอย่างน้อย 1 ไฟล์ภาพ, 1 ไฟล์เสียง, และ 1 ไฟล์วิดีโอ")

        # แยกไฟล์ตามประเภท
        image_file = audio_file = video_file = None
        for f in files:
            ext = os.path.splitext(f)[1].lower()
            if ext in ['.png', '.jpg', '.jpeg'] and not image_file:
                image_file = f
            elif ext in ['.wav', '.mp3', '.flac', '.ogg', '.aac'] and not audio_file:
                audio_file = f
            elif ext in ['.mp4', '.avi', '.mov', '.mkv'] and not video_file:
                video_file = f

        if not image_file:
            raise ValueError("ไม่พบไฟล์ภาพ (รองรับ .png, .jpg)")
        if not audio_file:
            raise ValueError("ไม่พบไฟล์เสียง (รองรับ .wav, .mp3)")
        if not video_file:
            raise ValueError("ไม่พบไฟล์วิดีโอ (รองรับ .mp4, .avi, .mov, .mkv)")

        try:
            print("[INFO] Encrypting and hiding data (Mode 3)...")

            # กำหนด key
            if config.get('random_aes_m3', False):
                key_str = self.gen_secure_key()
            else:
                password = config.get('aes_password_m3', '').strip()
                if not password:
                    raise ValueError("กรุณากรอกรหัสผ่าน AES หรือเลือก 'สุ่ม'")
                key_str = self.stretch_key(password)

            # เข้ารหัส
            iv_b64, ct_b64, _, encrypted_b64 = self.encrypt_aes(text, key_str)
            key_b64 = base64.b64encode(key_str.encode()).decode()
            self.result_display.append(f"<span style='color: #00d4ff;'>🔑 Key (string): {key_str}</span>")

            # แบ่งข้อมูลเป็น 3 ส่วน
            if config.get('custom_split_m3', False):
                image_ratio = config.get('image_ratio', 33)
                audio_ratio = config.get('audio_ratio', 33)
                video_ratio = config.get('video_ratio', 34)
                parts = self.split_msg_by_ratio(encrypted_b64, [image_ratio, audio_ratio, video_ratio])
            else:
                parts = self.split_msg(encrypted_b64, 3)

            part_img, part_audio, part_video = parts

            # 🔹 ซ่อนในภาพ
            out_img_name = f"stego_img_{uuid.uuid4().hex[:8]}.png"
            out_img_path = os.path.join(output_dir, out_img_name)
            success_img = self.hide_lsb_image(image_file, part_img, out_img_path)
            if not success_img:
                raise RuntimeError("ล้มเหลวในการซ่อนข้อมูลในภาพ")

            # 🔹 ซ่อนในเสียง
            out_audio_dir = os.path.join(output_dir, "audio_output_m3")
            os.makedirs(out_audio_dir, exist_ok=True)
            out_audio_path = self.hide_lsb_audio(audio_file, part_audio, out_audio_dir)
            if not out_audio_path:
                raise RuntimeError("ล้มเหลวในการซ่อนข้อมูลในเสียง")

            # 🔹 ซ่อนในวิดีโอ (ใช้ thread เหมือนเดิม)
            filename = os.path.splitext(os.path.basename(video_file))[0]
            timestamp = time.strftime("%Y%m%d%H%M%S")
            out_video_dir = os.path.join(output_dir, "video_output_m3")
            os.makedirs(out_video_dir, exist_ok=True)
            out_video_path = os.path.join(out_video_dir, f"{filename}_hidden_{timestamp}.avi")

            self.result_display.append("<font color='blue'>🎬 กำลังซ่อนข้อมูลในวิดีโอ...</font>")
            self.worker = VideoSteganographyWorker(
                video_file, out_video_path, part_video,
                self.extract_frames, self.encode_message_to_last_frame, self.combine_frames_to_video,
                self.hide_message_in_image, self.extract_message_from_image
            )
            self.worker.finished.connect(self.on_hide_finished_mode3)
            self.worker.start()

            # แสดงผลภาพและเสียงทันที (วิดีโอรอ thread เสร็จ)
            self.result_display.append(f"<span style='color: #00ff88;'>✅ ซ่อนข้อมูลในภาพและเสียงสำเร็จ!</span>")
            self.result_display.append(f"<span style='color: #00d4ff;'>🖼️ ภาพ: {out_img_name}</span>")
            self.result_display.append(f"<span style='color: #00d4ff;'>🎵 เสียง: {os.path.basename(out_audio_path)}</span>")

        except Exception as e:
            self.result_display.append(f"<span style='color: #ff4444;'>❌ ข้อผิดพลาด: {str(e)}</span>")
            raise











    def execute_mode4(self, item):
        """Mode 4: AES + RSA + Metadata"""
        # Placeholder: Implement AES + RSA with metadata embedding
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 4: รอการพัฒนา (AES + RSA + Metadata)</span>")
        # TODO: Encrypt text with AES, encrypt key with RSA, embed in metadata

    def execute_mode5(self, item):
        """Mode 5: GPG + Metadata + EOF"""
        # Placeholder: Implement GPG encryption with metadata and EOF embedding
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 5: รอการพัฒนา (GPG + Metadata + EOF)</span>")
        # TODO: Use GPG for encryption, embed in metadata and EOF

    def execute_mode6(self, item):
        """Mode 6: AES + LSB + Metadata + Checksum"""
        # Placeholder: Implement AES + LSB with metadata and checksum
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 6: รอการพัฒนา (AES + LSB + Metadata + Checksum)</span>")
        # TODO: Encrypt with AES, embed in LSB, store key in metadata, add checksum

    def execute_mode7(self, item):
        """Mode 7: Multi-layer Transformation"""
        # Placeholder: Implement multi-layer encoding
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 7: รอการพัฒนา (Multi-layer Transformation)</span>")
        # TODO: Apply base64, gzip, AES, embed in multiple locations

    def execute_mode8(self, item):
        """Mode 8: AES + GPG + Multi Media"""
        # Placeholder: Implement AES + GPG across multiple media
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 8: รอการพัฒนา (AES + GPG + Multi Media)</span>")
        # TODO: Encrypt with AES and GPG, distribute across media

    def execute_mode9(self, item):
        """Mode 9: Nested Stego"""
        # Placeholder: Implement nested steganography
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 9: รอการพัฒนา (Nested Stego)</span>")
        # TODO: Nest data in image, then audio, then video

    def execute_mode10(self, item):
        """Mode 10: Split + Layered + Time-lock"""
        # Placeholder: Implement split, layered, and time-locked steganography
        self.result_display.append("<span style='color: #ffeb3b;'>⚠️ Mode 10: รอการพัฒนา (Split + Layered + Time-lock)</span>")
        # TODO: Split data, apply layered encryption, add time-lock
        
        
        
        

    def on_extract_mode_changed(self, index):
        """Handle extract mode selection change"""
        mode_id = self.extract_mode_dropdown.itemData(index)
        self.current_extract_mode_id = mode_id
        self.update_extract_mode_description(mode_id)
        self.update_extract_mode_configuration(mode_id)

    def update_extract_mode_description(self, mode_id):
        """Update extract mode description based on selected mode"""
        descriptions = {
            1: "🔓 ถอดข้อความจากไฟล์ภาพและเสียง → รวมข้อมูล → ถอดรหัส AES\n🔑 ต้องมี: ไฟล์ภาพ + ไฟล์เสียง + รหัสผ่าน AES",
            2: "🔓 ถอดข้อมูลจาก metadata วิดีโอ → ถอดรหัส RSA → เปิดไฟล์ DOCX\n🔑 ต้องมี: ไฟล์วิดีโอ + RSA private key",
            3: "🔓 ถอดข้อมูลจากภาพ/เสียง/วิดีโอ → รวม 3 ส่วน → ถอดรหัส AES\n🔑 ต้องมี: ไฟล์ภาพ + เสียง + วิดีโอ + รหัสผ่าน AES",
            4: "🔓 ถอดข้อมูลจาก metadata → ถอดรหัส RSA key → ถอดรหัส AES\n🔑 ต้องมี: ไฟล์เสียง + วิดีโอ + RSA private key",
            5: "🔓 ถอดข้อมูลจาก metadata + EOF → ถอดรหัส GPG\n🔑 ต้องมี: ไฟล์ภาพ + เสียง + วิดีโอ + GPG private key",
            6: "🔓 ถอดข้อมูลจาก LSB → ถอด key จาก metadata → ตรวจสอบ checksum\n🔑 ต้องมี: ไฟล์ภาพ + เสียง + วิดีโอ + รหัสผ่าน AES",
            7: "🔓 ถอดข้อมูลหลายชั้น → ถอดรหัส AES → decompress → decode\n🔑 ต้องมี: ไฟล์ภาพ + วิดีโอ + เสียง + รหัสผ่าน AES + RSA private key",
            8: "🔓 ถอดข้อมูลจากหลาย media → ถอดรหัส GPG key → ถอดรหัส AES\n🔑 ต้องมี: ไฟล์วิดีโอ + ภาพ + เสียง + รหัสผ่าน AES + GPG private key",
            9: "🔓 ถอดข้อมูลแบบซ้อน → ถอดจากวิดีโอ → เสียง → ภาพ → ถอดรหัส AES\n🔑 ต้องมี: ไฟล์ภาพ + เสียง + วิดีโอ + รหัสผ่าน AES + RSA private key",
            10: "🔓 ถอดข้อมูลหลายส่วน → ถอดรหัส RSA keys → รวมข้อมูล → ตรวจสอบ time-lock\n🔑 ต้องมี: หลายไฟล์ + Master password + RSA private keys"
        }
        
        self.extract_mode_description.setText(descriptions.get(mode_id, "ไม่พบคำอธิบาย"))

    def update_extract_mode_configuration(self, mode_id):
        """Update extract configuration UI based on selected mode"""
        # Clear existing configuration completely
        self.clear_extract_config_layout()
        
        # Create configuration based on mode
        if mode_id == 1:
            self.create_extract_mode1_config()
        elif mode_id == 2:
            self.create_extract_mode2_config()
        elif mode_id == 3:
            self.create_extract_mode3_config()
        elif mode_id == 4:
            self.create_extract_mode4_config()
        elif mode_id == 5:
            self.create_extract_mode5_config()
        elif mode_id == 6:
            self.create_extract_mode6_config()
        elif mode_id == 7:
            self.create_extract_mode7_config()
        elif mode_id == 8:
            self.create_extract_mode8_config()
        elif mode_id == 9:
            self.create_extract_mode9_config()
        elif mode_id == 10:
            self.create_extract_mode10_config()

    def clear_extract_config_layout(self):
        """Completely clear the extract configuration layout"""
        def clear_layout(layout):
            if layout is not None:
                while layout.count():
                    item = layout.takeAt(0)
                    widget = item.widget()
                    if widget is not None:
                        widget.deleteLater()
                    else:
                        child_layout = item.layout()
                        if child_layout is not None:
                            clear_layout(child_layout)
        
        clear_layout(self.extract_config_layout)
        self.extract_config_group.update()

    def create_extract_mode1_config(self):
        """Extract Configuration for Mode 1: AES + Split Half"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password = QLineEdit()
        self.extract_aes_password.setPlaceholderText("กรอกรหัสผ่าน AES ที่ใช้ในการซ่อน")
        self.extract_aes_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password, 1, 1)
        
        show_password_btn = QPushButton("👁️ แสดงรหัสผ่าน")
        show_password_btn.setCheckable(True)
        show_password_btn.toggled.connect(lambda checked: self.extract_aes_password.setEchoMode(QLineEdit.Normal if checked else QLineEdit.Password))
        layout.addWidget(show_password_btn, 2, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 3, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (ที่มีข้อมูลครึ่งแรก)\n• 1 ไฟล์เสียง (ที่มีข้อมูลครึ่งหลัง)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 4, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)


    def create_extract_mode2_config(self):
        """Extract Configuration for Mode 2: DOCX + RSA + Video Metadata"""
        layout = QGridLayout()
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA สำหรับถอดรหัส:"), 0, 0, 1, 2)
        layout.addWidget(QLabel("RSA Private Key:"), 1, 0)
        
        rsa_layout = QHBoxLayout()
        self.extract_rsa_private_key = QLineEdit()
        self.extract_rsa_private_key.setPlaceholderText("เส้นทางไฟล์ private key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_rsa_private_key, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.extract_rsa_private_key)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 1, 1)
        
        layout.addWidget(QLabel("RSA Key Password:"), 2, 0)
        self.extract_rsa_password = QLineEdit()
        self.extract_rsa_password.setPlaceholderText("รหัสผ่าน private key (ถ้ามี)")
        self.extract_rsa_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_rsa_password, 2, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 3, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์วิดีโอ (ที่มี metadata ซ่อนอยู่)\n• 1 ไฟล์ DOCX (ที่ถูกป้องกันรหัสผ่าน)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 4, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)



    def create_extract_mode3_config(self):
        """Extract Configuration for Mode 3: AES + Split 3 Parts"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password_m3 = QLineEdit()
        self.extract_aes_password_m3.setPlaceholderText("กรอกรหัสผ่าน AES ที่ใช้ในการซ่อน")
        self.extract_aes_password_m3.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password_m3, 1, 1)
        
        # Split Configuration
        layout.addWidget(QLabel("✂️ การตั้งค่าการรวมข้อมูล:"), 2, 0, 1, 2)
        
        self.extract_auto_detect = QCheckBox("ตรวจจับสัดส่วนอัตโนมัติ")
        self.extract_auto_detect.setChecked(True)
        layout.addWidget(self.extract_auto_detect, 3, 0, 1, 2)
        
        # Manual split ratios
        split_ratio_layout = QHBoxLayout()
        split_ratio_layout.addWidget(QLabel("สัดส่วน - ภาพ:"))
        self.extract_image_ratio = QSpinBox()
        self.extract_image_ratio.setRange(1, 100)
        self.extract_image_ratio.setValue(33)
        self.extract_image_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.extract_image_ratio)
        
        split_ratio_layout.addWidget(QLabel("เสียง:"))
        self.extract_audio_ratio = QSpinBox()
        self.extract_audio_ratio.setRange(1, 100)
        self.extract_audio_ratio.setValue(33)
        self.extract_audio_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.extract_audio_ratio)
        
        split_ratio_layout.addWidget(QLabel("วิดีโอ:"))
        self.extract_video_ratio = QSpinBox()
        self.extract_video_ratio.setRange(1, 100)
        self.extract_video_ratio.setValue(34)
        self.extract_video_ratio.setEnabled(False)
        split_ratio_layout.addWidget(self.extract_video_ratio)
        
        self.extract_auto_detect.toggled.connect(lambda checked: [
            self.extract_image_ratio.setEnabled(not checked),
            self.extract_audio_ratio.setEnabled(not checked),
            self.extract_video_ratio.setEnabled(not checked)
        ])
        
        layout.addLayout(split_ratio_layout, 4, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (ส่วนที่ 1)\n• 1 ไฟล์เสียง (ส่วนที่ 2)\n• 1 ไฟล์วิดีโอ (ส่วนที่ 3)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode4_config(self):
        """Extract Configuration for Mode 4: AES + RSA + Metadata"""
        layout = QGridLayout()
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Private Key:"), 1, 0)
        rsa_layout = QHBoxLayout()
        self.extract_rsa_private_key_m4 = QLineEdit()
        self.extract_rsa_private_key_m4.setPlaceholderText("เส้นทางไฟล์ private key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_rsa_private_key_m4, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.extract_rsa_private_key_m4)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 1, 1)
        
        layout.addWidget(QLabel("RSA Key Password:"), 2, 0)
        self.extract_rsa_password_m4 = QLineEdit()
        self.extract_rsa_password_m4.setPlaceholderText("รหัสผ่าน private key (ถ้ามี)")
        self.extract_rsa_password_m4.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_rsa_password_m4, 2, 1)
        
        # Metadata Options
        layout.addWidget(QLabel("📋 ตัวเลือกการถอด Metadata:"), 3, 0, 1, 2)
        
        self.extract_verify_integrity = QCheckBox("ตรวจสอบความถูกต้องของข้อมูล")
        self.extract_verify_integrity.setChecked(True)
        layout.addWidget(self.extract_verify_integrity, 4, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์เสียง (มี ciphertext)\n• 1 ไฟล์วิดีโอ (มี RSA-encrypted key)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode5_config(self):
        """Extract Configuration for Mode 5: GPG + Metadata + EOF"""
        layout = QGridLayout()
        
        # GPG Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า GPG สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("GPG Private Key:"), 1, 0)
        gpg_layout = QHBoxLayout()
        self.extract_gpg_private_key = QLineEdit()
        self.extract_gpg_private_key.setPlaceholderText("เส้นทางไฟล์ GPG private key")
        gpg_browse_btn = QPushButton("เลือกไฟล์")
        gpg_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_gpg_private_key, "GPG Key Files (*.asc *.gpg)"))
        gpg_layout.addWidget(self.extract_gpg_private_key)
        gpg_layout.addWidget(gpg_browse_btn)
        layout.addLayout(gpg_layout, 1, 1)
        
        layout.addWidget(QLabel("GPG Passphrase:"), 2, 0)
        self.extract_gpg_passphrase = QLineEdit()
        self.extract_gpg_passphrase.setPlaceholderText("รหัสผ่าน GPG key")
        self.extract_gpg_passphrase.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_gpg_passphrase, 2, 1)
        
        # EOF Options
        layout.addWidget(QLabel("➕ ตัวเลือกการถอด EOF:"), 3, 0, 1, 2)
        
        self.extract_try_all_formats = QCheckBox("ลองทุกรูปแบบ EOF")
        self.extract_try_all_formats.setChecked(True)
        layout.addWidget(self.extract_try_all_formats, 4, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (metadata)\n• 1 ไฟล์เสียง (metadata)\n• 1 ไฟล์วิดีโอ (EOF data)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode6_config(self):
        """Extract Configuration for Mode 6: AES + LSB + Metadata + Checksum"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password_m6 = QLineEdit()
        self.extract_aes_password_m6.setPlaceholderText("กรอกรหัสผ่าน AES (ถ้าทราบ)")
        self.extract_aes_password_m6.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password_m6, 1, 1)
        
        # LSB Configuration
        layout.addWidget(QLabel("🖼️ การตั้งค่าการถอด LSB:"), 2, 0, 1, 2)
        
        self.extract_auto_detect_lsb = QCheckBox("ตรวจจับการตั้งค่า LSB อัตโนมัติ")
        self.extract_auto_detect_lsb.setChecked(True)
        layout.addWidget(self.extract_auto_detect_lsb, 3, 0, 1, 2)
        
        lsb_manual_layout = QHBoxLayout()
        lsb_manual_layout.addWidget(QLabel("LSB bits:"))
        self.extract_lsb_bits = QSpinBox()
        self.extract_lsb_bits.setRange(1, 4)
        self.extract_lsb_bits.setValue(2)
        self.extract_lsb_bits.setEnabled(False)
        lsb_manual_layout.addWidget(self.extract_lsb_bits)
        
        lsb_manual_layout.addWidget(QLabel("ช่องสี:"))
        self.extract_lsb_channels = QComboBox()
        self.extract_lsb_channels.addItems(["RGB ทั้งหมด", "เฉพาะ Red", "เฉพาะ Green", "เฉพาะ Blue"])
        self.extract_lsb_channels.setEnabled(False)
        lsb_manual_layout.addWidget(self.extract_lsb_channels)
        
        self.extract_auto_detect_lsb.toggled.connect(lambda checked: [
            self.extract_lsb_bits.setEnabled(not checked),
            self.extract_lsb_channels.setEnabled(not checked)
        ])
        
        layout.addLayout(lsb_manual_layout, 4, 0, 1, 2)
        
        # Checksum Verification
        layout.addWidget(QLabel("🧪 การตรวจสอบ Checksum:"), 5, 0, 1, 2)
        
        self.extract_verify_checksum = QCheckBox("ตรวจสอบ checksum อัตโนมัติ")
        self.extract_verify_checksum.setChecked(True)
        layout.addWidget(self.extract_verify_checksum, 6, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 7, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (LSB data)\n• 1 ไฟล์เสียง (AES key metadata)\n• 1 ไฟล์วิดีโอ (checksum metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 8, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode7_config(self):
        """Extract Configuration for Mode 7: Multi-layer Transformation"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password_m7 = QLineEdit()
        self.extract_aes_password_m7.setPlaceholderText("กรอกรหัสผ่าน AES")
        self.extract_aes_password_m7.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password_m7, 1, 1)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA (สำหรับ key backup):"), 2, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Private Key:"), 3, 0)
        rsa_layout = QHBoxLayout()
        self.extract_rsa_private_key_m7 = QLineEdit()
        self.extract_rsa_private_key_m7.setPlaceholderText("เส้นทางไฟล์ private key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_rsa_private_key_m7, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.extract_rsa_private_key_m7)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 3, 1)
        
        # Transformation Options
        layout.addWidget(QLabel("🧪 ตัวเลือกการถอดแปลง:"), 4, 0, 1, 2)
        
        self.extract_auto_detect_transform = QCheckBox("ตรวจจับการแปลงอัตโนมัติ")
        self.extract_auto_detect_transform.setChecked(True)
        layout.addWidget(self.extract_auto_detect_transform, 5, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 6, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (LSB data)\n• 1 ไฟล์วิดีโอ (metadata)\n• 1 ไฟล์เสียง (EOF backup)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 7, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode8_config(self):
        """Extract Configuration for Mode 8: AES + GPG + Multi Media"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password_m8 = QLineEdit()
        self.extract_aes_password_m8.setPlaceholderText("กรอกรหัสผ่าน AES")
        self.extract_aes_password_m8.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password_m8, 1, 1)
        
        # GPG Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า GPG สำหรับถอดรหัส:"), 2, 0, 1, 2)
        
        layout.addWidget(QLabel("GPG Private Key:"), 3, 0)
        gpg_layout = QHBoxLayout()
        self.extract_gpg_private_key_m8 = QLineEdit()
        self.extract_gpg_private_key_m8.setPlaceholderText("เส้นทางไฟล์ GPG private key")
        gpg_browse_btn = QPushButton("เลือกไฟล์")
        gpg_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_gpg_private_key_m8, "GPG Key Files (*.asc *.gpg)"))
        gpg_layout.addWidget(self.extract_gpg_private_key_m8)
        gpg_layout.addWidget(gpg_browse_btn)
        layout.addLayout(gpg_layout, 3, 1)
        
        layout.addWidget(QLabel("GPG Passphrase:"), 4, 0)
        self.extract_gpg_passphrase_m8 = QLineEdit()
        self.extract_gpg_passphrase_m8.setPlaceholderText("รหัสผ่าน GPG key")
        self.extract_gpg_passphrase_m8.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_gpg_passphrase_m8, 4, 1)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 5, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์วิดีโอ (ciphertext metadata)\n• 1 ไฟล์ภาพ (GPG-encrypted key)\n• 1 ไฟล์เสียง (hash EOF)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 6, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode9_config(self):
        """Extract Configuration for Mode 9: Nested Stego"""
        layout = QGridLayout()
        
        # AES Configuration
        layout.addWidget(QLabel("🔐 การตั้งค่า AES สำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("รหัสผ่าน AES:"), 1, 0)
        self.extract_aes_password_m9 = QLineEdit()
        self.extract_aes_password_m9.setPlaceholderText("กรอกรหัสผ่าน AES")
        self.extract_aes_password_m9.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_aes_password_m9, 1, 1)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA สำหรับถอดรหัส:"), 2, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Private Key:"), 3, 0)
        rsa_layout = QHBoxLayout()
        self.extract_rsa_private_key_m9 = QLineEdit()
        self.extract_rsa_private_key_m9.setPlaceholderText("เส้นทางไฟล์ private key")
        rsa_browse_btn = QPushButton("เลือกไฟล์")
        rsa_browse_btn.clicked.connect(lambda: self.browse_file(self.extract_rsa_private_key_m9, "RSA Key Files (*.pem *.key)"))
        rsa_layout.addWidget(self.extract_rsa_private_key_m9)
        rsa_layout.addWidget(rsa_browse_btn)
        layout.addLayout(rsa_layout, 3, 1)
        
        # Nesting Configuration
        layout.addWidget(QLabel("🌀 การตั้งค่าการถอดแบบซ้อน:"), 4, 0, 1, 2)
        
        self.extract_auto_detect_nesting = QCheckBox("ตรวจจับการซ้อนอัตโนมัติ")
        self.extract_auto_detect_nesting.setChecked(True)
        layout.addWidget(self.extract_auto_detect_nesting, 5, 0, 1, 2)
        
        nesting_manual_layout = QHBoxLayout()
        nesting_manual_layout.addWidget(QLabel("จำนวนชั้น:"))
        self.extract_nesting_levels = QSpinBox()
        self.extract_nesting_levels.setRange(2, 5)
        self.extract_nesting_levels.setValue(3)
        self.extract_nesting_levels.setEnabled(False)
        nesting_manual_layout.addWidget(self.extract_nesting_levels)
        
        nesting_manual_layout.addWidget(QLabel("Encoding:"))
        self.extract_encoding_method = QComboBox()
        self.extract_encoding_method.addItems(["Base64", "Hex", "Binary", "Custom"])
        self.extract_encoding_method.setEnabled(False)
        nesting_manual_layout.addWidget(self.extract_encoding_method)
        
        self.extract_auto_detect_nesting.toggled.connect(lambda checked: [
            self.extract_nesting_levels.setEnabled(not checked),
            self.extract_encoding_method.setEnabled(not checked)
        ])
        
        layout.addLayout(nesting_manual_layout, 6, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 7, 0, 1, 2)
        req_label = QLabel("• 1 ไฟล์ภาพ (ชั้นใน)\n• 1 ไฟล์เสียง (ชั้นกลาง)\n• 1 ไฟล์วิดีโอ (ชั้นนอก + key metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 8, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def create_extract_mode10_config(self):
        """Extract Configuration for Mode 10: Split + Layered + Time-lock"""
        layout = QGridLayout()
        
        # Master Password
        layout.addWidget(QLabel("🔐 การตั้งค่าหลักสำหรับถอดรหัส:"), 0, 0, 1, 2)
        
        layout.addWidget(QLabel("Master Password:"), 1, 0)
        self.extract_master_password = QLineEdit()
        self.extract_master_password.setPlaceholderText("รหัสผ่านหลักที่ใช้ในการซ่อน")
        self.extract_master_password.setEchoMode(QLineEdit.Password)
        layout.addWidget(self.extract_master_password, 1, 1)
        
        # RSA Configuration
        layout.addWidget(QLabel("🔑 การตั้งค่า RSA Keys:"), 2, 0, 1, 2)
        
        layout.addWidget(QLabel("RSA Private Keys:"), 3, 0)
        rsa_keys_layout = QVBoxLayout()
        
        self.extract_rsa_key_list = QListWidget()
        self.extract_rsa_key_list.setMaximumHeight(100)
        rsa_keys_layout.addWidget(self.extract_rsa_key_list)
        
        rsa_buttons_layout = QHBoxLayout()
        add_extract_rsa_btn = QPushButton("เพิ่ม RSA Private Key")
        add_extract_rsa_btn.clicked.connect(self.add_extract_rsa_key)
        remove_extract_rsa_btn = QPushButton("ลบ RSA Key")
        remove_extract_rsa_btn.setObjectName("clearButton")
        remove_extract_rsa_btn.clicked.connect(self.remove_extract_rsa_key)
        rsa_buttons_layout.addWidget(add_extract_rsa_btn)
        rsa_buttons_layout.addWidget(remove_extract_rsa_btn)
        rsa_keys_layout.addLayout(rsa_buttons_layout)
        
        layout.addLayout(rsa_keys_layout, 3, 1)
        
        # Time-lock Configuration
        layout.addWidget(QLabel("⏳ การตั้งค่า Time-lock:"), 4, 0, 1, 2)
        
        self.extract_ignore_timelock = QCheckBox("ข้าม Time-lock (ถ้าเป็นเจ้าของ)")
        layout.addWidget(self.extract_ignore_timelock, 5, 0, 1, 2)
        
        # Auto-detect Configuration
        layout.addWidget(QLabel("🔍 การตรวจจับอัตโนมัติ:"), 6, 0, 1, 2)
        
        self.extract_auto_detect_parts = QCheckBox("ตรวจจับจำนวนส่วนอัตโนมัติ")
        self.extract_auto_detect_parts.setChecked(True)
        layout.addWidget(self.extract_auto_detect_parts, 7, 0, 1, 2)
        
        # File Requirements
        layout.addWidget(QLabel("📋 ไฟล์ที่ต้องการ:"), 8, 0, 1, 2)
        req_label = QLabel("• หลายไฟล์ภาพ/เสียง/วิดีโอ (ตามจำนวนส่วนที่แบ่ง)\n• หลาย RSA private keys\n• 1 ไฟล์ ZIP (metadata)")
        req_label.setObjectName("warningLabel")
        layout.addWidget(req_label, 9, 0, 1, 2)
        
        self.extract_config_layout.addLayout(layout)

    def add_extract_rsa_key(self):
        """Add RSA private key to the extract list"""
        file_path, _ = QFileDialog.getOpenFileName(self, "เลือก RSA Private Key", "", "RSA Key Files (*.pem *.key);;All Files (*.*)")
        if file_path:
            self.extract_rsa_key_list.addItem(os.path.basename(file_path) + " → " + file_path)

    def remove_extract_rsa_key(self):
        """Remove selected RSA key from the extract list"""
        current_row = self.extract_rsa_key_list.currentRow()
        if current_row >= 0:
            self.extract_rsa_key_list.takeItem(current_row)

    def select_extract_files(self):
        """Select files for extraction"""
        file_types = "All Supported Files (*.png *.jpg *.jpeg *.bmp *.gif *.tiff *.wav *.mp3 *.flac *.ogg *.mp4 *.avi *.mov *.mkv *.docx *.pdf);;Images (*.png *.jpg *.jpeg *.bmp *.gif *.tiff);;Audio (*.wav *.mp3 *.flac *.ogg);;Video (*.mp4 *.avi *.mov *.mkv);;Documents (*.docx *.pdf);;All Files (*.*)"
        files, _ = QFileDialog.getOpenFileNames(self, "เลือกไฟล์สำหรับถอดข้อมูล (Select Files for Extraction)", "", file_types)
        if files:
            self.selected_extract_files.extend(files)
            self.update_extract_files_table()

    def update_extract_files_table(self):
        """Update the extract files table display"""
        self.extract_files_table.setRowCount(len(self.selected_extract_files))
        for row, file_path in enumerate(self.selected_extract_files):
            file_name = os.path.basename(file_path)
            self.extract_files_table.setItem(row, 0, QTableWidgetItem(file_name))
            
            file_ext = os.path.splitext(file_path)[1].lower()
            file_type = "ไม่ระบุ (Unknown)"
            if file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif', '.tiff']:
                file_type = "ภาพ (Image)"
            elif file_ext in ['.wav', '.mp3', '.flac', '.ogg']:
                file_type = "เสียง (Audio)"
            elif file_ext in ['.mp4', '.avi', '.mov', '.mkv']:
                file_type = "วิดีโอ (Video)"
            elif file_ext in ['.docx', '.pdf']:
                file_type = "เอกสาร (Document)"
            self.extract_files_table.setItem(row, 1, QTableWidgetItem(file_type))
            
            try:
                size = os.path.getsize(file_path)
                size_str = self.format_size(size)
            except:
                size_str = "ไม่ทราบ"
            self.extract_files_table.setItem(row, 2, QTableWidgetItem(size_str))
            
            # Status column
            status = "พร้อมถอด (Ready to Extract)"
            self.extract_files_table.setItem(row, 3, QTableWidgetItem(status))

    def clear_extract_files(self):
        """Clear selected extract files"""
        self.selected_extract_files.clear()
        self.extract_files_table.setRowCount(0)
        self.result_display.append("<span style='color: #00d4ff;'>🗑️ รายการไฟล์ถอดข้อมูลถูกล้างแล้ว (Extract file list cleared).</span>")

    def copy_extracted_text(self):
        """Copy extracted text to clipboard"""
        text = self.extracted_text_display.toPlainText()
        if text.strip():
            from PyQt5.QtWidgets import QApplication
            QApplication.clipboard().setText(text)
            self.result_display.append("<span style='color: #00ff88;'>✅ คัดลอกข้อความไปยัง clipboard แล้ว (Text copied to clipboard).</span>")
        else:
            QMessageBox.warning(self, "คำเตือน", "ไม่มีข้อความให้คัดลอก")

    def save_extracted_text(self):
        """Save extracted text to file"""
        text = self.extracted_text_display.toPlainText()
        if not text.strip():
            QMessageBox.warning(self, "คำเตือน", "ไม่มีข้อความให้บันทึก")
            return
            
        file_path, _ = QFileDialog.getSaveFileName(self, "บันทึกข้อความที่ถอดได้", "", "Text Files (*.txt);;All Files (*.*)")
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as file:
                    file.write(text)
                    self.result_display.append(f"<span style='color: #00ff88;'>✅ บันทึกข้อความที่ถอดได้ไปยังไฟล์: {os.path.basename(file_path)}</span>")
            except Exception as e:
                QMessageBox.warning(self, "ข้อผิดพลาด", f"ไม่สามารถบันทึกไฟล์ได้: {str(e)}")



# -------------------------------- fun ten mode ---------------------------------------

    def str_to_bin(self,text):
        try:
            return ''.join(format(b, '08b') for b in text.encode('utf-8'))
        except:
            return ""

    def bin_to_str(self,bin_str):
        try:
            if bin_str.endswith("00000000"):
                bin_str = bin_str[:-8]
            n = len(bin_str)
            byte_data = int(bin_str, 2).to_bytes((n + 7) // 8, 'big')
            return byte_data.decode('utf-8')
        except:
            return ""

    def gen_secure_key(self, length=32):
        """Generate cryptographically secure key"""
        chars = string.ascii_letters + string.digits
        return ''.join(random.SystemRandom().choice(chars) for _ in range(length))

    def stretch_key(self, password, length=32):
        """Stretch password to 32-byte key using simple method (for demo). Use PBKDF2 in production."""
        key = (password * (length // len(password) + 1))[:length]
        return key.ljust(length, 'X')  # ทำให้ได้ 32 ตัว

    def encrypt_aes(self, text, key_str=None):
        if not key_str:
            key_str = self.gen_secure_key()
        key = key_str.encode('utf-8')
        cipher = AES.new(key, AES.MODE_CBC)
        ct = cipher.encrypt(pad(text.encode('utf-8'), AES.block_size))
        iv_b64 = base64.b64encode(cipher.iv).decode()
        ct_b64 = base64.b64encode(ct).decode()
        encrypted_b64 = base64.b64encode(cipher.iv + ct).decode()
        return iv_b64, ct_b64, key_str, encrypted_b64

    def split_msg(self, msg, parts=2):
        length = len(msg)
        return [msg[i*length//parts : (i+1)*length//parts] for i in range(parts)]

    def hide_lsb_image(self, img_path, msg, out_path):
        try:
            img = Image.open(img_path)
            if img.mode != 'RGB':
                if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                    bg = Image.new("RGB", img.size, (255, 255, 255))
                    if img.mode == 'P':
                        img = img.convert("RGBA")
                    alpha = img.split()[-1]
                    bg.paste(img.convert('RGB'), mask=alpha)
                    img = bg
                else:
                    img = img.convert('RGB')
            arr = np.array(img)
            bin_msg = ''.join(format(b, '08b') for b in msg.encode('utf-8')) + '00000000'
            h, w, c = arr.shape
            if len(bin_msg) > h * w * c:
                raise ValueError("ข้อความยาวเกินไปสำหรับภาพนี้")
            idx = 0
            for i in range(h):
                for j in range(w):
                    for k in range(c):
                        if idx < len(bin_msg):
                            arr[i, j, k] = (arr[i, j, k] & 0xFE) | int(bin_msg[idx])
                            idx += 1
            Image.fromarray(arr).save(out_path, 'PNG')
            return True
        except Exception as e:
            print(f"[ERROR] hide_lsb_image: {e}")
            return False

    def hide_lsb_audio(self, audio_path, data, out_dir):
        if not os.path.exists(audio_path) or not data.strip():
            return None
        ext = os.path.splitext(audio_path)[1].lower()
        temp_wav = None
        use_path = audio_path
        if ext != ".wav":
            audio = AudioSegment.from_file(audio_path)
            temp_wav = f"temp_{uuid.uuid4().hex}.wav"
            audio.export(temp_wav, format="wav")
            use_path = temp_wav
        try:
            with wave.open(use_path, 'rb') as af:
                params = af.getparams()
                frames = af.readframes(af.getnframes())
            audio_data = np.frombuffer(frames, dtype=np.uint8)
            bin_data = ''.join(format(b, '08b') for b in data.encode('utf-8')) + '00000000'
            if len(bin_data) > len(audio_data):
                raise ValueError("ข้อมูลยาวเกินไปสำหรับไฟล์เสียง")
            mod_data = audio_data.copy()
            for i, bit in enumerate(bin_data):
                mod_data[i] = (mod_data[i] & 0xFE) | int(bit)
            os.makedirs(out_dir, exist_ok=True)
            out_path = os.path.join(out_dir, os.path.splitext(os.path.basename(audio_path))[0] + "_hidden.wav")
            with wave.open(out_path, 'wb') as of:
                of.setparams(params)
                of.writeframes(mod_data.tobytes())
            return out_path
        finally:
            if temp_wav and os.path.exists(temp_wav):
                os.remove(temp_wav)


    def extract_lsb_image(self, img_path):
        """Extract binary data from image using LSB, stop at terminator"""
        try:
            img = Image.open(img_path).convert('RGB')
            arr = np.array(img)
            bin_data = ""
            for val in arr.flatten():
                bin_data += str(val & 1)
                if len(bin_data) % 8 == 0 and bin_data.endswith("00000000"):
                    # แปลงเป็น string
                    byte_val = int(bin_data[:-8], 2).to_bytes((len(bin_data)-8 + 7)//8, 'big')
                    return byte_val.decode('utf-8', errors='replace')
            # หากไม่เจอ terminator
            if len(bin_data) == 0:
                return ""
            n = (len(bin_data) + 7) // 8
            byte_val = int(bin_data, 2).to_bytes(n, 'big')
            return byte_val.decode('utf-8', errors='replace')
        except Exception as e:
            print(f"[ERROR] extract_lsb_image: {e}")
            return ""

    def extract_lsb_audio(self, audio_path):
        """Extract binary data from audio using LSB"""
        try:
            if not os.path.exists(audio_path):
                return ""
            with wave.open(audio_path, 'rb') as af:
                frames = af.readframes(af.getnframes())
            audio_data = np.frombuffer(frames, dtype=np.uint8)
            bin_data = ""
            for b in audio_data:
                bin_data += str(b & 1)
                if len(bin_data) % 8 == 0 and bin_data.endswith("00000000"):
                    byte_val = int(bin_data[:-8], 2).to_bytes((len(bin_data)-8 + 7)//8, 'big')
                    return byte_val.decode('utf-8', errors='replace')
            if len(bin_data) == 0:
                return ""
            n = (len(bin_data) + 7) // 8
            byte_val = int(bin_data, 2).to_bytes(n, 'big')
            return byte_val.decode('utf-8', errors='replace')
        except Exception as e:
            print(f"[ERROR] extract_lsb_audio: {e}")
            return ""



# -------------------------------- fun ten mode ---------------------------------------






    def execute_extraction(self):
        """Extract hidden data from selected files based on current mode"""
        if not self.selected_extract_files:
            QMessageBox.warning(self, "คำเตือน", "กรุณาเลือกไฟล์เพื่อถอดข้อมูล")
            return

        self.progress_bar.setVisible(True)
        self.extracted_text_display.clear()
        self.result_display.append("🔍 เริ่มกระบวนการถอดข้อมูล...")

        try:
            mode_id = self.current_extract_mode_id
            files = self.selected_extract_files.copy()

            
            if mode_id == 1:
                self.extract_mode1(files)
            elif mode_id == 2:
                self.extract_mode2(files)
            elif mode_id == 3:
                self.extract_mode3(files)
            elif mode_id == 4:
                self.extract_mode4(files)
            elif mode_id == 5:
                self.extract_mode5(files)
            elif mode_id == 6:
                self.extract_mode6(files)
            elif mode_id == 7:
                self.extract_mode7(files)
            elif mode_id == 8:
                self.extract_mode8(files)
            elif mode_id == 9:
                self.extract_mode9(files)
            elif mode_id == 10:
                self.extract_mode10(files)
            else:
                self.result_display.append("<span style='color: #ffeb3b;'>⚠️ ไม่รองรับโหมดนี้</span>")

        except Exception as e:
            self.result_display.append(f"<span style='color: #ff4444;'>❌ ข้อผิดพลาด: {str(e)}</span>")
        finally:
            self.progress_bar.setVisible(False)

    def extract_mode1(self, files):
        """Mode 1: Extract data from image and audio → Combine → Decrypt with AES"""
        try:
            # --- 1. ค้นหาไฟล์ภาพและเสียง ---
            image_file = None
            audio_file = None
            for f in files:
                ext = os.path.splitext(f)[1].lower()
                if ext in ['.png', '.jpg', '.jpeg'] and not image_file:
                    image_file = f
                elif ext in ['.wav', '.mp3', '.flac', '.ogg'] and not audio_file:
                    audio_file = f

            if not image_file:
                raise ValueError("ไม่พบไฟล์ภาพ (.png/.jpg)")
            if not audio_file:
                raise ValueError("ไม่พบไฟล์เสียง (.wav/.mp3)")

            self.result_display.append(f"<span style='color: #00d4ff;'>🔍 ใช้ภาพ: {os.path.basename(image_file)}</span>")
            self.result_display.append(f"<span style='color: #00d4ff;'>🎵 ใช้เสียง: {os.path.basename(audio_file)}</span>")

            # --- 2. ดึงข้อมูลจากภาพ (LSB) ---
            p1 = self.extract_lsb_image(image_file)
            if not p1:
                raise ValueError("ไม่สามารถดึงข้อมูลจากภาพได้ (อาจไม่มีข้อมูลซ่อนอยู่)")
            self.result_display.append(f"<span style='color: #ffeb3b;'>🖼️ ดึงข้อมูลจากภาพสำเร็จ (ความยาว: {len(p1)})</span>")

            # --- 3. ดึงข้อมูลจากเสียง (LSB) ---
            p2 = self.extract_lsb_audio(audio_file)
            if not p2:
                raise ValueError("ไม่สามารถดึงข้อมูลจากเสียงได้")
            self.result_display.append(f"<span style='color: #ffeb3b;'>🔊 ดึงข้อมูลจากเสียงสำเร็จ (ความยาว: {len(p2)})</span>")

            # --- 4. รวมข้อมูล ---
            combined_b64 = p1 + p2
            self.result_display.append(f"<span style='color: #00d4ff;'>🔗 รวมข้อมูล: {len(combined_b64)} ตัวอักษร</span>")

            # ตรวจสอบว่าเป็น Base64 ที่ถูกต้องหรือไม่
            try:
                encrypted_bytes = base64.b64decode(combined_b64, validate=True)
                if len(encrypted_bytes) < 16:
                    raise ValueError("ข้อมูลสั้นเกินไป (ต้องมี IV + ciphertext)")
            except Exception as e:
                raise ValueError(f"Base64 ไม่ถูกต้อง: {e}")

            # --- 5. ดึง key จากหน้าจอ ---
            key_input = self.extract_aes_password.text().strip()
            if not key_input:
                raise ValueError("กรุณากรอก Key (Base64 หรือ 32-character string)")

            # ถ้าเป็น Base64 → ถอดออกมา
            try:
                if len(key_input) % 4 == 0 and all(c in "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/=" for c in key_input):
                    key_str = base64.b64decode(key_input).decode('utf-8')
                else:
                    key_str = key_input  # สมมติว่าเป็นข้อความธรรมดา
            except Exception:
                key_str = key_input  # ใช้ raw string เลย

            if len(key_str) != 32:
                raise ValueError(f"Key ต้องมี 32 ตัวอักษร (ตอนนี้: {len(key_str)})")

            # --- 6. ถอดรหัส AES-CBC ---
            try:
                iv = encrypted_bytes[:16]
                ct = encrypted_bytes[16:]
                cipher = AES.new(key_str.encode('utf-8'), AES.MODE_CBC, iv)
                decrypted_padded = cipher.decrypt(ct)
                decrypted_text = unpad(decrypted_padded, AES.block_size).decode('utf-8')

                # --- 7. แสดงผลลัพธ์ ---
                self.extracted_text_display.setPlainText(decrypted_text)
                self.result_display.append("<span style='color: #00ff88;'>✅ ถอดรหัสสำเร็จ!</span>")
                self.result_display.append(f"<span style='color: #00d4ff;'>📝 ข้อความ: {len(decrypted_text)} ตัวอักษร</span>")
                return decrypted_text

            except Exception as e:
                if "Padding" in str(e):
                    raise ValueError("ถอดรหัสไม่สำเร็จ: Padding error (key หรือ IV ผิด)")
                else:
                    raise ValueError(f"ถอดรหัสไม่สำเร็จ: {str(e)}")

        except Exception as e:
            self.result_display.append(f"<span style='color: #ff4444;'>❌ ข้อผิดพลาด: {str(e)}</span>")
            return None

    def extract_mode2(self, files):
        """Extract DOCX + RSA + Video Metadata"""
        try:
            self.result_display.append("🔓 เริ่มกระบวนการถอดรหัสและดึงข้อมูล...")
            
            # ดึงค่าจาก UI
            private_key_path = self.extract_rsa_private_key.text()
            protected_docx = next((f for f in files if f.endswith('.docx')), None)
            video_file = next((f for f in files if os.path.splitext(f)[1].lower() in ['.mp4', '.avi', '.mov', '.mkv']), None)
            
            # ตรวจสอบไฟล์
            if not video_file:
                raise ValueError("ไม่พบไฟล์วิดีโอสำหรับถอดข้อมูล")
            if not protected_docx:
                raise ValueError("ไม่พบไฟล์ DOCX ที่ถูกป้องกันรหัสผ่าน")
            if not os.path.exists(private_key_path):
                raise FileNotFoundError(f"ไม่พบไฟล์ private key: {private_key_path}")
            
            # ดึง metadata จากวิดีโอ
            result = subprocess.run([
                "ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", video_file
            ], capture_output=True, text=True, check=True)
            
            metadata = json.loads(result.stdout)
            encoded_encrypted_password = metadata['format']['tags'].get('comment')
            
            if not encoded_encrypted_password:
                raise ValueError("ไม่พบ metadata ที่ชื่อ 'comment'")
            
            self.result_display.append(f"🔍 ดึงข้อมูล metadata สำเร็จ: {encoded_encrypted_password}")
            
            # โหลด private key
            with open(private_key_path, "rb") as key_file:
                private_key = serialization.load_pem_private_key(
                    key_file.read(),
                    password=self.extract_rsa_password.text().encode() if self.extract_rsa_password.text() else None
                )
            
            # ถอดรหัสรหัสผ่าน
            encrypted_password = base64.b64decode(encoded_encrypted_password)
            decrypted_password = private_key.decrypt(
                encrypted_password,
                padding.OAEP(
                    mgf=padding.MGF1(algorithm=hashes.SHA256()),
                    algorithm=hashes.SHA256(),
                    label=None
                )
            )
            docx_password = decrypted_password.decode()
            self.result_display.append(f"🔓 ถอดรหัสรหัสผ่านสำเร็จ: {docx_password}")
            
            # ถอดรหัส DOCX
            decrypted_docx = os.path.join(self.output_dir, "decrypted.docx")
            
            with open(protected_docx, "rb") as f_in:
                with open(decrypted_docx, "wb") as f_out:
                    office_file = msoffcrypto.OfficeFile(f_in)
                    office_file.load_key(password=docx_password)
                    office_file.decrypt(f_out)
            
            self.result_display.append(f"📄 ถอดรหัส DOCX เรียบร้อย: {decrypted_docx}")
            
            # อ่านเนื้อหา
            doc = Document(decrypted_docx)
            content = "\n".join([p.text for p in doc.paragraphs])
            self.extracted_text_display.setPlainText(content)
            
        except Exception as e:
            self.result_display.append(f"❌ ข้อผิดพลาด: {str(e)}")
            self.progress_bar.setValue(0)
        else:
            self.progress_bar.setValue(100)
            self.result_display.append("✅ กระบวนการสำเร็จสมบูรณ์!")


















